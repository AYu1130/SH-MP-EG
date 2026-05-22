#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
添加第3-6章、参考文献（不含16、17、18）、致谢
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_chinese_font(run, font_name='宋体', size=12, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold

def add_heading_custom(doc, text, level=1):
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    
    if level == 1:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(16)
        run.font.bold = True
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        heading.paragraph_format.space_before = Pt(24)
        heading.paragraph_format.space_after = Pt(18)
    elif level == 2:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(14)
        run.font.bold = True
        heading.paragraph_format.space_before = Pt(18)
        heading.paragraph_format.space_after = Pt(12)
    else:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(12)
        run.font.bold = True
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)
    
    heading.paragraph_format.line_spacing = 1.5
    return heading

def add_paragraph_custom(doc, text, first_line_indent=True, fontsize=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(fontsize)
    
    p.paragraph_format.line_spacing = 1.5
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(6)
    
    return p

def add_figure_placeholder(doc, caption, note=None):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(f'【{caption} - 请在此处插入图片】')
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(11)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(f'图{caption}')
    set_chinese_font(run, '宋体', 10.5)
    p.paragraph_format.space_after = Pt(12)
    
    if note:
        p = doc.add_paragraph()
        run = p.add_run(f'注：{note}')
        set_chinese_font(run, '宋体', 10)
        p.paragraph_format.space_after = Pt(6)

def add_table_from_data(doc, headers, rows, caption=None):
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(f'表{caption}')
        set_chinese_font(run, '宋体', 10.5, True)
        p.paragraph_format.space_after = Pt(6)
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(10)
                run.font.bold = True
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
            for p in row_cells[i].paragraphs:
                for run in p.runs:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(10)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    return table

def add_code_block(doc, code_lines, caption=None):
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(caption)
        set_chinese_font(run, '宋体', 10.5, True)
        p.paragraph_format.space_after = Pt(6)
    
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F2F2F2')
    shading_elm.set(qn('w:val'), 'clear')
    cell._element.get_or_add_tcPr().append(shading_elm)
    
    for line in code_lines:
        p = cell.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.3)
    
    doc.add_paragraph()

# 打开文档
doc = Document('e:\\100_study\\120_Project\\CapstoneProject\\SmartHome_MultiProtocol_EdgeIntelligent_Gateway\\SH-MP-EG\\docs\\output\\面向智能家居的多协议边缘智能网关设计_去除16_17_18_部分2.docx')

# ========== 第3章 ==========
doc.add_page_break()
add_heading_custom(doc, '第3章 多协议边缘网关总体方案与开发平台', level=1)

add_heading_custom(doc, '3.1 网关总体方案设计', level=2)
add_heading_custom(doc, '3.1.1 系统整体分层架构', level=3)

para_311 = [
    '本系统采用分层架构设计思想，将多协议边缘网关划分为设备接入层、协议适配层、消息路由与处理层、应用与联动管理层四个功能层次[1,8,10]。各层之间通过标准化接口进行交互，实现功能解耦和模块独立。',
    
    '设备接入层位于架构最底层，负责与各类智能家居终端设备进行物理连接和原始数据收发。该层包含Wi-Fi接入模块、BLE接入模块以及预留的其他协议扩展接口[5,6]。Wi-Fi接入模块通过TCP/UDP Socket与ESP32-S3设备建立通信连接，接收传感器上报的环境数据；BLE接入模块通过蓝牙协议栈与STM32F103设备进行GATT特征值读写，获取传感器数据或下发控制指令[4,13]。',
    
    '协议适配层承担异构协议数据的解析与转换职能。该层维护一组协议适配器（Protocol Adapter），每种支持的协议对应一个适配器实例[1,3]。Wi-Fi适配器负责解析ESP32通过Socket发送的原始数据帧，提取有效载荷并进行格式校验；BLE适配器负责处理蓝牙GATT通信，将特征值数据转换为内部数据结构；MQTT适配器作为输出侧适配器，将内部数据结构序列化为JSON格式的MQTT消息[8,14]。',
    
    '消息路由与处理层是网关的数据中枢，负责消息的接收、缓存、路由和转发。该层部署EMQX消息代理作为核心组件，处理MQTT消息的发布/订阅逻辑[10,14]。同时，该层包含消息路由器（Message Router），根据预定义的规则将消息分发至不同的处理终端：本地联动引擎、Web管理后台或云端转发模块。消息队列机制确保高并发场景下的数据处理可靠性[1,10]。',
    
    '应用与联动管理层面向终端用户和上层应用提供功能接口。Node-RED可视化编程环境作为本地联动引擎，允许用户通过拖拽方式配置自动化规则[8,11]；Flask Web框架构建的管理后台提供设备管理、状态监控、系统配置等交互功能；云端桥接模块在需要时将数据转发至远程服务器，实现远程访问和大数据分析[5,6]。'
]

for text in para_311:
    add_paragraph_custom(doc, text)

# 添加系统数据流图占位符
add_figure_placeholder(
    doc,
    '3-1 系统数据流架构图',
    '展示终端设备→网关→协议转换→MQTT Broker→各应用层的完整数据流向'
)

add_heading_custom(doc, '3.1.2 系统数据流架构', level=3)
add_paragraph_custom(doc, '系统的数据流向如下：终端设备（Wi-Fi/BLE）采集环境数据后，通过各自的无线协议发送至网关[1,5]。网关的协议适配层接收原始数据，进行协议解析和格式转换，生成统一的JSON格式内部消息。内部消息经消息路由器分发，一部分进入Node-RED规则引擎触发本地联动逻辑，一部分存储至SQLite本地数据库，另一部分通过MQTT Broker发布至Web管理后台供用户查看[8,10]。')

add_paragraph_custom(doc, '在网络连通状态下，部分数据还可经云端桥接模块上传至远程服务器，实现跨网络访问。控制指令的数据流方向相反：用户通过Web界面或联动规则触发控制指令，指令经MQTT Broker路由至对应协议的适配器，转换为目标协议格式后下发至终端设备执行[1,4,14]。')

# 添加数据流时序图占位符
add_figure_placeholder(
    doc,
    '3-2 数据流向时序图',
    '展示传感器上报→网关处理→规则引擎执行→控制指令下发的完整时序'
)

# 3.2 硬件平台
add_heading_custom(doc, '3.2 硬件平台设计与实现', level=2)
add_heading_custom(doc, '3.2.1 核心主控硬件选型：树莓派4B', level=3)

para_321 = [
    '树莓派4B（Raspberry Pi 4 Model B）作为网关的核心主控单元，承担数据处理、协议转换、规则执行和资源管理等核心任务[1,10]。选型树莓派4B的主要依据包括：',
    
    '计算性能方面，树莓派4B搭载Broadcom BCM2711四核Cortex-A72处理器，主频1.5GHz，相比上一代产品性能提升显著[5]。4GB RAM配置足以支撑Python运行环境、EMQX消息代理、Node-RED规则引擎以及Flask Web服务的并发运行，为多任务处理提供充足的内存资源[8]。',
    
    '接口扩展方面，树莓派4B提供40针GPIO接口、4个USB 2.0接口、千兆以太网接口以及双频Wi-Fi和蓝牙5.0无线连接能力[1]。GPIO接口可用于连接外设扩展板，USB接口可接入Zigbee协调器、BLE适配器等扩展模块，为系统预留了充足的硬件扩展空间[10]。',
    
    '软件生态方面，Raspberry Pi OS基于Debian Linux发行版，拥有丰富的软件包资源和活跃的社区支持[6]。Python 3、Node.js、SQLite等本系统依赖的运行环境和数据库均提供官方支持，部署和配置过程简便可靠[8]。',
    
    '成本效益方面，树莓派4B的零售价格约300-400元人民币，相比工控计算机或专用网关设备具有明显的价格优势，符合消费级智能家居产品的成本约束[1,5]。'
]

for text in para_321:
    add_paragraph_custom(doc, text)

add_heading_custom(doc, '3.2.2 Wi-Fi终端节点：ESP32-S3', level=3)
add_paragraph_custom(doc, 'ESP32-S3作为Wi-Fi终端节点的核心处理器，负责环境数据采集和无线数据传输[5,6]。该芯片由乐鑫科技推出，具有以下技术特点：无线连接能力方面，ESP32-S3集成2.4GHz Wi-Fi和Bluetooth 5（LE）双模无线通信能力，在本系统中主要使用其Wi-Fi功能与网关建立网络连接[1,10]。Wi-Fi支持802.11 b/g/n标准，最大传输速率150Mbps，信号覆盖范围和穿透能力在住宅环境中表现良好。')

add_paragraph_custom(doc, '计算能力方面，ESP32-S3搭载Xtensa® 32位LX7双核处理器，主频最高240MHz，内置512KB SRAM和384KB ROM，可支持复杂的传感器数据处理和边缘推理任务[5,13]。芯片还集成了向量指令集，为轻量级AI应用提供硬件加速。外设接口方面，ESP32-S3提供丰富的外设接口，包括SPI、I2C、I2S、UART、ADC、DAC等，便于连接各类传感器和执行器[6]。在本系统中，ESP32-S3通过I2C接口连接SHT30温湿度传感器和BH1750光照传感器，定时采集环境数据并通过Wi-Fi上报至网关[1,10]。')

add_heading_custom(doc, '3.2.3 BLE终端节点：STM32F103', level=3)
add_paragraph_custom(doc, 'STM32F103系列微控制器作为BLE终端节点的核心，配合外部BLE模块实现蓝牙设备的接入[3,4,13]。选型依据如下：低功耗特性方面，STM32F103基于ARM Cortex-M3内核，运行功耗低至数毫安级别，待机功耗可降至微安级，适合电池供电的长期运行场景[13]。配合BLE模块的超低功耗设计，可实现纽扣电池数年的续航时间。')

add_paragraph_custom(doc, '实时响应方面，STM32F103支持多种低功耗模式和中断唤醒机制，能够在接收到网关查询请求时快速从睡眠状态恢复，完成数据上报后迅速返回低功耗状态，最小化能量消耗[4,6]。开发便利性方面，STM32拥有完善的开发工具链和丰富的外设库支持，开发者可通过HAL库或标准外设库快速完成硬件接口配置和功能实现[3,13]。成熟的社区生态为开发过程中的问题排查提供了丰富的参考资源。在本系统中，STM32F103通过UART接口与BLE模块（如HC-08、CC2541等）进行通信，实现蓝牙协议的底层交互[1,5]。')

# 添加硬件实物图占位符
add_figure_placeholder(
    doc,
    '3-3 硬件系统实物图',
    '展示树莓派4B网关、ESP32-S3 Wi-Fi节点、STM32F103 BLE节点的实物连接照片'
)

# 添加硬件连接框图占位符
add_figure_placeholder(
    doc,
    '3-4 硬件系统连接框图',
    '展示各硬件模块之间的物理连接关系和数据流向'
)

add_heading_custom(doc, '3.2.4 硬件系统实物图', level=3)
add_paragraph_custom(doc, '【请在此处插入硬件系统实物照片，包含树莓派4B作为主控网关、ESP32-S3开发板作为Wi-Fi传感器节点、STM32F103开发板作为BLE传感器节点，以及传感器模块和执行器模块的整体连接照片。照片应清晰展示各模块之间的物理连接关系和接线方式。】')

# 3.3 软件平台
add_heading_custom(doc, '3.3 软件平台搭建与运行环境配置', level=2)

para_33 = [
    '网关软件平台建立在Raspberry Pi OS（64位版本）之上[6,8]。该操作系统基于Debian Bullseye发行版，针对树莓派硬件进行了深度优化，支持全部硬件接口并提供稳定的运行环境。系统安装采用官方Imager工具完成镜像烧录，首次启动后通过raspi-config工具进行基础配置，包括启用SSH远程访问、配置Wi-Fi网络连接、设置主机名和时区等[1,5]。为保障系统安全性，修改默认pi用户密码，创建专用的服务运行账户，限制关键进程的权限范围。',
    
    'EMQX是一款高性能的开源MQTT消息代理，支持MQTT 3.1/3.1.1/5.0协议标准，具备百万级并发连接和消息吞吐能力[10,14]。在本系统中，EMQX承担设备消息汇聚、路由和转发的核心职能。安装过程通过官方APT仓库完成，配置优化针对家庭网关场景进行调整：修改emqx.conf配置文件，将监听端口限制在本地回环地址以确保安全性；调整消息队列大小和会话过期时间，优化内存占用；启用WebSocket监听器，支持浏览器端MQTT客户端直接连接[8,10]。',
    
    'Node-RED是基于Node.js的可视化编程工具，通过拖拽节点和连接连线的方式实现应用程序的快速构建[1,8,11]。在本系统中，Node-RED作为本地联动引擎，负责自动化规则的执行。联动规则示例：实现"光照自动控灯"功能。配置mqtt-in节点订阅"home/+/light_sensor/+/data"主题，接收所有光照传感器数据；function节点解析消息中的亮度值，当检测到数值低于阈值时，构造开灯指令消息；mqtt-out节点将指令发送至"home/+/light/+/control"主题[10,12]。',
    
    '网关的协议适配层和Web管理后台基于Python 3.9开发[5,8]。通过pip包管理器安装依赖库：flask、paho-mqtt、bleak、sqlite3等。Flask作为轻量级Web框架，用于构建设备管理后台的RESTful API服务。应用采用蓝图（Blueprint）模式组织代码结构，将设备管理、用户认证、系统配置等功能模块化[1,10]。协议适配程序采用多线程架构，主线程负责系统初始化和协调管理，独立的工作线程分别处理Wi-Fi Socket通信、BLE蓝牙扫描和MQTT消息收发，通过线程安全的队列实现数据交换[8,14]。'
]

for text in para_33:
    add_paragraph_custom(doc, text)

# 添加软件架构图占位符
add_figure_placeholder(
    doc,
    '3-5 软件系统架构图',
    '展示应用服务层、业务逻辑层、协议适配层、硬件抽象层的层次结构和模块关系'
)

print("第3章完成")

# ========== 第4章 ==========
doc.add_page_break()
add_heading_custom(doc, '第4章 多协议边缘网关软件系统设计与实现', level=1)

add_heading_custom(doc, '4.1 软件系统总体设计', level=2)
add_heading_custom(doc, '4.1.1 模块化设计思想', level=3)

para_411 = [
    '本系统软件设计遵循模块化、分层化的架构原则，将复杂的多协议网关功能分解为若干高内聚、低耦合的功能模块[1,8,10]。每个模块承担明确的职责，通过定义良好的接口进行交互，便于独立开发、测试和维护。软件架构自上而下划分为四个层次：应用服务层、业务逻辑层、协议适配层和硬件抽象层[5,6]。',
    
    '应用服务层面向用户提供交互接口，包括Web管理后台和RESTful API；业务逻辑层实现设备管理、联动规则、数据缓存等核心功能；协议适配层处理异构协议的解析与转换；硬件抽象层封装底层硬件接口，向上层提供统一的操作接口[1,14]。模块间的依赖关系遵循单向原则：上层模块可调用下层模块的接口，但下层模块不依赖上层模块的具体实现。这种设计保证了核心功能（如协议适配）的稳定性，同时允许上层应用（如Web界面）独立演进[8,10]。'
]

for text in para_411:
    add_paragraph_custom(doc, text)

# 表4-1
add_table_from_data(
    doc,
    ['层次', '模块名称', '功能描述', '技术实现'],
    [
        ['应用服务层', 'Web管理台', '提供可视化设备管理界面', 'Flask + HTML/JS'],
        ['应用服务层', 'REST API', '提供外部系统对接接口', 'Flask-RESTful'],
        ['业务逻辑层', '设备管理器', '设备注册、发现、状态维护', 'Python Class'],
        ['业务逻辑层', '联动引擎', '本地自动化规则执行', 'Node-RED'],
        ['业务逻辑层', '数据缓存', '离线数据存储与补传', 'SQLite'],
        ['协议适配层', 'Wi-Fi适配器', 'TCP/UDP Socket数据收发', 'Python Socket'],
        ['协议适配层', 'BLE适配器', '蓝牙GATT通信管理', 'Bleak库'],
        ['协议适配层', 'MQTT适配器', 'MQTT消息发布/订阅', 'Paho-MQTT'],
        ['协议适配层', '协议转换器', '数据格式转换与封装', 'JSON Schema'],
        ['硬件抽象层', 'GPIO控制', '树莓派GPIO接口操作', 'RPi.GPIO'],
        ['硬件抽象层', '串口通信', 'UART设备通信', 'PySerial'],
    ],
    caption='4-1 软件模块层次结构'
)

add_heading_custom(doc, '4.1.2 模块层次结构', level=3)
add_paragraph_custom(doc, '模块间的依赖关系遵循单向原则：上层模块可调用下层模块的接口，但下层模块不依赖上层模块的具体实现[1,8]。这种设计保证了核心功能（如协议适配）的稳定性，同时允许上层应用（如Web界面）独立演进。')

# 添加模块层次图占位符
add_figure_placeholder(
    doc,
    '4-1 模块层次结构图',
    '展示应用服务层→业务逻辑层→协议适配层→硬件抽象层的层次关系和模块组成'
)

# 4.2
add_heading_custom(doc, '4.2 统一数据传输模型与主题规范设计', level=2)
add_heading_custom(doc, '4.2.1 通用JSON数据格式定义', level=3)

para_42 = [
    '为实现异构协议数据的标准化表示，本系统定义了一套通用的JSON数据格式[1,8,10]。所有设备上报的数据和网关转发的消息均采用此格式封装，确保数据处理逻辑的一致性和可扩展性。',
    
    '温度传感器数据示例：{"device_id": "wifi_temp_001", "device_type": "temperature", "protocol": "wifi", "timestamp": 1704067200, "location": "livingroom", "data": {"temperature": 25.6, "humidity": 58.2, "unit": "celsius"}, "status": "online"}[10,14]。',
    
    '光照传感器数据示例：{"device_id": "ble_light_002", "device_type": "light", "protocol": "ble", "timestamp": 1704067200, "location": "bedroom", "data": {"illuminance": 350, "unit": "lux"}, "status": "online", "battery": 85}[8,13]。'
]

for text in para_42:
    add_paragraph_custom(doc, text)

# 表4-2
add_table_from_data(
    doc,
    ['字段名', '数据类型', '必填', '说明'],
    [
        ['device_id', 'String', '是', '设备唯一标识符'],
        ['device_type', 'String', '是', '设备类型'],
        ['protocol', 'String', '是', '接入协议：wifi/ble/zigbee'],
        ['timestamp', 'Integer', '是', '数据采集时间戳（Unix）'],
        ['location', 'String', '否', '设备位置信息'],
        ['data', 'Object', '是', '传感器数据对象'],
        ['status', 'String', '是', '设备状态：online/offline/error'],
        ['battery', 'Integer', '否', '电池电量百分比（0-100）'],
    ],
    caption='4-2 通用JSON数据格式字段定义'
)

add_heading_custom(doc, '4.2.2 MQTT主题命名规范', level=3)
add_paragraph_custom(doc, '本系统遵循层次化的MQTT主题命名规范，采用"home/区域/设备类型/设备ID/数据类型"的五级结构[10,12,14]。主题层级定义：第一级（home）根主题，标识智能家居系统；第二级（区域）房间或功能区域；第三级（设备类型）设备功能分类；第四级（设备ID）设备唯一标识符；第五级（数据类型）消息类型。QoS等级选择策略：设备数据上报使用QoS 1，控制指令使用QoS 1，心跳消息使用QoS 0[6,14]。')

# 添加主题层级图占位符
add_figure_placeholder(
    doc,
    '4-2 MQTT主题层级结构图',
    '展示home/区域/设备类型/设备ID/数据类型的五级主题结构及通配符使用示例'
)

# 4.3
add_heading_custom(doc, '4.3 协议适配与数据收发模块设计实现', level=2)
add_heading_custom(doc, '4.3.1 Wi-Fi终端接入模块', level=3)

para_431 = [
    'Wi-Fi终端接入模块负责管理ESP32-S3设备通过Wi-Fi网络与网关的通信连接[1,5,6]。模块采用TCP Socket Server模式运行，监听指定端口（默认8888），等待设备连接建立。程序逻辑架构：模块由主监听线程和多个设备处理线程组成。主监听线程持续监听端口，当检测到新的设备连接请求时，验证设备身份（通过预置的设备ID白名单），身份验证通过后创建独立的设备处理线程，专门负责与该设备的后续通信[8,14]。',
    
    '数据流处理流程：设备建立TCP连接后，首先发送身份注册帧，包含设备ID、设备类型等信息；网关验证通过后，进入数据通信阶段[1,10]。设备周期性地发送传感器数据帧，帧格式采用简单的文本协议：设备ID、数据类型、数值、时间戳以逗号分隔，以换行符结束。网关接收到数据帧后，解析字段内容，填充至预定义的JSON模板，生成统一格式的消息[8,10]。'
]

for text in para_431:
    add_paragraph_custom(doc, text)

# 添加Wi-Fi适配器流程图占位符
add_figure_placeholder(
    doc,
    '4-3 Wi-Fi终端接入模块流程图',
    '展示Socket监听→设备连接→身份验证→数据接收→协议解析→JSON封装的完整流程'
)

# Wi-Fi适配器代码
wifi_code = [
    'class WiFiAdapter:',
    '    def __init__(self, host=\'0.0.0.0\', port=8888):',
    '        self.host = host',
    '        self.port = port',
    '        self.devices = {}',
    '        self.server_socket = socket.socket(',
    '            socket.AF_INET, socket.SOCK_STREAM)',
    '        self.server_socket.setsockopt(',
    '            socket.SOL_SOCKET, socket.SO_REUSEADDR, 1)',
    '    ',
    '    def start(self):',
    '        self.server_socket.bind((self.host, self.port))',
    '        self.server_socket.listen(5)',
    '        threading.Thread(target=self._accept_loop, ',
    '                        daemon=True).start()',
    '    ',
    '    def _accept_loop(self):',
    '        while True:',
    '            client_socket, address = self.server_socket.accept()',
    '            device_thread = threading.Thread(',
    '                target=self._handle_device,',
    '                args=(client_socket, address),',
    '                daemon=True)',
    '            device_thread.start()',
    '    ',
    '    def _handle_device(self, client_socket, address):',
    '        device_id = self._authenticate(client_socket)',
    '        if not device_id:',
    '            client_socket.close()',
    '            return',
    '        self.devices[device_id] = client_socket',
    '        while True:',
    '            try:',
    '                data = client_socket.recv(1024)',
    '                if not data:',
    '                    break',
    '                parsed_data = self._parse_wifi_frame(data)',
    '                unified_msg = self._convert_to_unified_format(',
    '                    parsed_data)',
    '                self._publish_to_mqtt(unified_msg)',
    '            except Exception as e:',
    '                logger.error(f"Device {device_id} error: {e}")',
    '                break',
    '        del self.devices[device_id]',
    '        client_socket.close()'
]

add_code_block(doc, wifi_code, '代码4-1 Wi-Fi适配器核心类实现')

add_heading_custom(doc, '4.3.2 BLE终端接入模块', level=3)
add_paragraph_custom(doc, 'BLE终端接入模块负责扫描、发现、连接BLE设备，并进行GATT特征值的读写操作[3,4,13]。模块基于Bleak库（Bluetooth Low Energy platform Agnostic Klient）实现，该库提供了跨平台的Python异步BLE接口。模块架构设计采用异步编程模型，充分利用Python asyncio库的高效并发处理能力[8,14]。')

add_paragraph_custom(doc, '核心组件包括：扫描器（Scanner）定期扫描周围BLE广播设备；连接器（Connector）管理与目标设备的GATT连接；通知处理器（Notification Handler）处理设备主动推送的特征值变更通知；指令发送器（Command Sender）向设备下发控制指令[1,5,6]。GATT通信协议：BLE设备采用UART over BLE模式进行数据传输，使用Nordic UART Service（NUS）UUID定义服务[4,13]。')

# 添加BLE适配器流程图占位符
add_figure_placeholder(
    doc,
    '4-4 BLE终端接入模块流程图',
    '展示蓝牙扫描→设备发现→GATT连接→特征值订阅→数据通知处理的流程'
)

# BLE适配器代码
ble_code = [
    'class BLEAdapter:',
    '    def __init__(self):',
    '        self.devices = {}',
    '        self.connected_devices = {}',
    '        self.scanner = BleakScanner()',
    '    ',
    '    async def start_scan(self):',
    '        self.scanner.register_detection_callback(',
    '            self._on_device_detected)',
    '        await self.scanner.start()',
    '    ',
    '    def _on_device_detected(self, device, advertisement_data):',
    '        device_id = self._extract_device_id(advertisement_data)',
    '        if device_id and device_id.startswith(\'ble_\'):',
    '            self.devices[device_id] = device.address',
    '            asyncio.create_task(self._connect_device(',
    '                device.address, device_id))',
    '    ',
    '    async def _connect_device(self, address, device_id):',
    '        client = BleakClient(address)',
    '        try:',
    '            await client.connect()',
    '            self.connected_devices[device_id] = client',
    '            await client.start_notify(',
    '                UART_TX_CHAR_UUID,',
    '                lambda s, d: self._on_notification(device_id, s, d))',
    '        except Exception as e:',
    '            logger.error(f"BLE connect error: {e}")',
    '    ',
    '    def _on_notification(self, device_id, sender, data):',
    '        parsed_data = self._parse_ble_payload(data)',
    '        unified_msg = self._convert_to_unified_format(',
    '            device_id, parsed_data)',
    '        self._publish_to_mqtt(unified_msg)'
]

add_code_block(doc, ble_code, '代码4-2 BLE适配器核心类实现')

print("第4章前半部分完成")

# 保存
doc.save('e:\\100_study\\120_Project\\CapstoneProject\\SmartHome_MultiProtocol_EdgeIntelligent_Gateway\\SH-MP-EG\\docs\\output\\面向智能家居的多协议边缘智能网关设计_去除16_17_18_部分3.docx')
print("部分3保存完成")
