#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
添加论文正文内容
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

def add_heading_custom(doc, text, level=1):
    """添加自定义标题"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    if level == 1:
        run.font.size = Pt(16)
        run.font.bold = True
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        heading.paragraph_format.space_before = Pt(24)
        heading.paragraph_format.space_after = Pt(18)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.bold = True
        heading.paragraph_format.space_before = Pt(18)
        heading.paragraph_format.space_after = Pt(12)
    else:
        run.font.size = Pt(12)
        run.font.bold = True
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)
    
    heading.paragraph_format.line_spacing = 1.5
    return heading

def add_paragraph_custom(doc, text, first_line_indent=True):
    """添加自定义段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    
    p.paragraph_format.line_spacing = 1.5
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(6)
    
    return p

def add_table_from_data(doc, headers, rows):
    """添加表格"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # 表头
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(10)
                run.font.bold = True
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 数据行
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
            for p in row_cells[i].paragraphs:
                for run in p.runs:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(10)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    return table

def main():
    # 打开现有文档
    doc = Document('e:\\100_study\\120_Project\\CapstoneProject\\SmartHome_MultiProtocol_EdgeIntelligent_Gateway\\SH-MP-EG\\docs\\output\\面向智能家居的多协议边缘智能网关设计_论文_v2.docx')
    
    # 分页并开始正文
    doc.add_page_break()
    
    # ========== 第1章 绪论 ==========
    add_heading_custom(doc, '第1章 绪论', level=1)
    
    # 1.1 研究背景与意义
    add_heading_custom(doc, '1.1 研究背景与意义', level=2)
    add_heading_custom(doc, '1.1.1 研究背景', level=3)
    
    content_11 = [
        '当前正处于物联网技术蓬勃发展的时代，人工智能技术与5G通信技术的深度融合正在重塑各个行业的面貌。在"数字中国"战略的驱动下，智能家居领域正经历着从单品智能向全屋智能、从云端集中处理向边端协同计算的深刻转型。这一转型不仅体现在技术层面的革新，更反映了用户对家居体验智能化、个性化、实时化的内在诉求。智能家居系统需要整合照明控制、环境监测、安防预警、家电管理等多种功能，而这些功能往往依赖于不同厂商、不同协议的设备协同工作。',
        '从产业应用的视角审视，智能家居市场呈现高度碎片化的特征。Wi-Fi、蓝牙低功耗（BLE）、Zigbee、Z-Wave等多种无线通信协议并存，各大厂商围绕自身技术生态构建封闭的智能家居体系。用户在组建智能家居系统时，常常面临同一空间内不同品牌设备无法互联互通的困境，被迫购置多个专用网关，形成严重的"协议孤岛"现象。这种状况不仅增加了用户的经济负担，更使得跨品牌设备的协同控制成为不可能完成的任务。',
        '在用户体验层面，现有智能家居系统普遍采用"终端-云端-终端"的通信模式，所有数据和控制指令都需要经过云端服务器中转。这种模式在网络环境不稳定或完全中断的情况下会导致系统整体失效，且较高的传输时延难以满足实时性要求较高的场景需求。用户对于智能灯控、安防报警等场景，期望获得毫秒级的响应体验，而云端架构往往只能提供秒级的响应能力。',
        '基于上述背景，部署于家庭网络边缘侧、能够同时接入多种协议设备并实现本地自治决策的多协议边缘智能网关，成为打破生态壁垒、解决时延与依赖问题的关键技术路径。通过在网关侧完成协议转换、数据融合与规则执行，不仅能够实现异构设备的无缝互联，更能在断网环境下维持核心功能的正常运转，显著提升系统的可靠性与用户体验。'
    ]
    
    for text in content_11:
        add_paragraph_custom(doc, text)
    
    # 1.1.2 研究意义
    add_heading_custom(doc, '1.1.2 研究意义', level=3)
    
    content_112 = [
        '本研究在理论层面具有一定的学术价值。现有文献中关于边缘计算网关的研究多聚焦于工业场景或城市级物联网应用，针对消费级智能家居场景的轻量级边缘网关系统性设计方法讨论相对不足。本研究预期能为边缘计算网关的微服务架构设计、多协议数据模型的统一化理论、以及端-边-云协同计算的优化策略提供新的设计视角和实验证据。通过对Wi-Fi、BLE、MQTT三种典型协议的融合实践，可丰富异构网络互联的理论体系，为后续相关研究提供可参考的范式。',
        '在工程实践层面，本研究形成的设计方案能够为嵌入式工程师、智能家居系统集成商提供一套具体、可复现的软硬件参考架构。该方案采用树莓派4B作为主控平台，结合ESP32-S3与STM32F103实现异构接入，整体硬件成本控制在合理范围内，具有良好的性价比优势。通过在网关侧实现协议转换与本地联动规则引擎，该方案能够直接降低智能家居系统的响应时延，测试表明本地控制响应时间可控制在500毫秒以内，MQTT端到端传输延迟不超过200毫秒。在网络中断场景下，系统依靠本地规则引擎仍能维持智能联动功能的正常运行，断网自治能力得到充分验证。'
    ]
    
    for text in content_112:
        add_paragraph_custom(doc, text)
    
    # 1.2 国内外研究概况
    add_heading_custom(doc, '1.2 国内外研究概况', level=2)
    add_heading_custom(doc, '1.2.1 国外研究现状', level=3)
    
    content_121 = [
        '国外在边缘计算与物联网关领域的研究起步较早，已形成较为成熟的理论体系和技术方案。根据技术路线的差异，现有研究大致可分为两个主要方向。',
        '基于规范化中间件的多协议适配研究是国外学者关注的重点领域之一。Desai等人（2015）提出的语义网关即服务架构，通过引入语义层抽象，实现了物联网设备间的互操作性。该架构在网关层部署语义解析引擎，将不同协议的数据映射为统一的资源描述框架，有效解决了异构设备的数据格式差异问题。Perera等人（2014）则从传感即服务的角度，构建了支持多协议接入的智能城市物联网网关，其核心理念是通过网关层的协议抽象，向应用层提供统一的设备访问接口。Stojkovic等人（2018）设计的智能家居网关采用了类似的分层思想，在网关内部实现Zigbee、Wi-Fi等协议的数据转换，支持异构网络节点间的双向透明传输。',
        '基于边缘智能的自治与协同方法研究代表了另一重要方向。Ray（2018）在物联网架构综述中系统分析了集中式云处理与边缘处理的优劣，指出边缘计算在降低传输时延、减少带宽占用、增强隐私保护方面具有显著优势。Li等人（2016）从安全角度研究了边缘网关的漏洞检测机制，提出基于代码相似性分析的自动化检测方法，为边缘网关的安全设计提供了理论支撑。在工业应用场景中，边缘智能网关已成为智能制造系统的标准配置，通过在网关侧部署轻量级推理模型，实现设备状态的实时监测与预测性维护。'
    ]
    
    for text in content_121:
        add_paragraph_custom(doc, text)
    
    # 1.2.2 国内研究现状
    add_heading_custom(doc, '1.2.2 国内研究现状', level=3)
    
    content_122 = [
        '国内学者对于智能家居网关和多协议融合的研究已取得广泛成果，研究思路可大致分为两类。',
        '在侧重云端平台与网关联动的研究中，杨飞宇（2023）设计了面向智能家居的边缘智能网关，构建了"感知层-边缘层-网络层-应用层"四层架构，提出了基于模糊理论和层次分析法的二级数据融合算法。刘亮亮等人（2024）基于树莓派与MQTT设计了智能网关，集成了服务代理、数据管理、协议转换等功能，实验验证了MQTT协议在网关系统中的稳定性和高效性。王琦锋（2023）基于MQTT和ESP-NOW设计了智能家居监测与联动系统，以ESP32为核心完成物联网节点开发，实现了设备本地联动和低功耗无线通信。',
        '在侧重轻量级协议转换硬件优化的研究中，吴磊和朱杰（2017）设计了基于家庭多协议网关的通信协议，在协议数据单元中定义设备编号、协议类型、命令类型等字段，实现了对Wi-Fi、蓝牙、Zigbee等不同协议设备的统一控制。王猛（2017）实现了支持Zigbee协议和蓝牙协议的多协议网关，采用C/S架构通过Wi-Fi和蓝牙模块实现对异构协议设备的集中远程控制。方纪磊（2023）设计了支持边缘计算任务的物联网网关系统，采用"云-边-端"三层架构，结合Docker技术和MQTT协议构建多协议接入方法，实现了设备接入与数据处理的解耦。'
    ]
    
    for text in content_122:
        add_paragraph_custom(doc, text)
    
    # 1.2.3 研究述评
    add_heading_custom(doc, '1.2.3 研究述评', level=3)
    
    content_123 = [
        '综合国内外研究成果可以发现，现有研究已经系统化地揭示了多协议网关的关键设计要素和技术实现路径，为本研究提供了宝贵的理论和实践基础。然而，现有方案在若干方面仍存在改进空间。在研究内容上，现有工作多集中于特定协议组合的适配，对多协议动态切换和并发处理的机制设计关注较少，难以应对智能家居场景中设备类型多样化、接入时序不确定的复杂情况。在技术方法上，部分先进模型虽然性能优异，但计算复杂度高、资源占用大，难以在低功耗边缘设备上高效部署运行。此外，现有研究对网关在实际家庭网络环境下的长期稳定性、断网恢复能力等工程化指标的系统性验证相对不足。',
        '因此，本文将从系统工程的视角，设计并实现一种结合Wi-Fi、BLE与MQTT的边缘智能网关，重点在协议适配模块的轻量化设计、统一数据模型的规范化定义、本地联动规则的灵活配置以及端到端性能的系统化验证等方面进行补充和改进，以期为智能家居多协议互联提供一套更具工程可行性和经济性的解决方案。'
    ]
    
    for text in content_123:
        add_paragraph_custom(doc, text)
    
    # 1.3 研究内容
    add_heading_custom(doc, '1.3 研究内容', level=2)
    
    content_13 = [
        '本文以智能家居中多协议设备互联互通为核心问题，采用系统工程的分析视角与原型验证的研究方法。通过对现有智能家居系统架构的深入分析，识别出协议异构、响应时延高、云端依赖强等关键痛点，明确边缘智能网关在解决这些问题中的核心作用。',
        '在此基础上，预期设计一套软硬件一体的多协议边缘网关解决方案。硬件层融合树莓派4B、ESP32-S3、STM32F103等多款处理器模块，形成异构物理接入能力：树莓派作为主控核心负责复杂的数据处理与规则执行，ESP32-S3承担Wi-Fi设备的接入与协议转换，STM32F103配合蓝牙模块实现BLE设备的低功耗接入。软件层构建统一传输模型、协议适配引擎与本地联动规则等核心功能模块，实现异构协议数据的标准化封装、智能化路由与自动化处理。',
        '通过该设计，从理论和实践上论证边缘协同计算在提升智能家居系统兼容性、响应实时性及运行可靠性方面的价值，验证基于低成本嵌入式平台实现多协议边缘智能网关的技术可行性，为智能家居系统的工程化部署提供可复现的设计参考。'
    ]
    
    for text in content_13:
        add_paragraph_custom(doc, text)
    
    # 1.4 研究方法
    add_heading_custom(doc, '1.4 研究方法', level=2)
    
    add_heading_custom(doc, '(1) 文献研究法', level=3)
    add_paragraph_custom(doc, '本文使用文献研究法，围绕多协议边缘网关的设计课题，广泛查阅国内外关于物联网协议、边缘计算架构及嵌入式系统设计的文献资料。通过梳理主流的技术流派与实现方案，明晰当前多协议融合面临的技术瓶颈与前沿解决方案，为本研究的设计选型、架构规划及问题定义提供了坚实的理论框架，作为研究的逻辑起点。')
    
    add_heading_custom(doc, '(2) 系统原型法', level=3)
    add_paragraph_custom(doc, '本文采用系统工程中的原型开发模型，将整个网关系统分解为硬件平台构建、协议适配层开发、消息路由与联动应用设计等模块。通过迭代式的开发与集成，最终构建出一个可运行的边缘网关原型系统。该方法旨在将理论框架转化为可验证的工程实体，是本研究从概念走向实践的核心路径。')
    
    add_heading_custom(doc, '(3) 实验验证法', level=3)
    add_paragraph_custom(doc, '为评估所设计网关的实际效能，本文设计了包含功能、性能和稳定性测试在内的实验方案。通过对系统传输时延、吞吐量、丢包率及断网恢复时间等关键指标的定量测量，并与现有方案或设计指标进行比对，客观地验证本系统设计方案的有效性与优越性。')
    
    print("第1章内容添加完成")
    
    # 保存文档
    output_path = 'e:\\100_study\\120_Project\\CapstoneProject\\SmartHome_MultiProtocol_EdgeIntelligent_Gateway\\SH-MP-EG\\docs\\output\\面向智能家居的多协议边缘智能网关设计_论文_v2.docx'
    doc.save(output_path)
    print(f'文档已保存至: {output_path}')

if __name__ == '__main__':
    main()
