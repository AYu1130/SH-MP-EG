#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成长沙学院毕业论文Word文档
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
import re

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_elm = parse_xml(r'<w:%s w:val="%s" w:sz="%s" w:space="0" w:color="%s" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>' % (
                edge, kwargs[edge]["val"], kwargs[edge]["sz"], kwargs[edge]["color"]))
            tcPr.append(edge_elm)

def add_heading_custom(doc, text, level=1, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT):
    """添加自定义标题"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '黑体')
    
    if level == 1:
        run.font.size = Pt(16)
        run.font.bold = True
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif level == 2:
        run.font.size = Pt(14)
        run.font.bold = True
    else:
        run.font.size = Pt(12)
        run.font.bold = True
    
    heading.paragraph_format.line_spacing = 1.5
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)
    
    return heading

def add_paragraph_custom(doc, text, first_line_indent=True, bold=False):
    """添加自定义段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    run.font.bold = bold
    
    p.paragraph_format.line_spacing = 1.5
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)  # 两个字符
    p.paragraph_format.space_after = Pt(6)
    
    return p

def add_table_custom(doc, headers, rows):
    """添加表格"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # 添加表头
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '宋体')
                run.font.size = Pt(10.5)
                run.font.bold = True
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 添加数据行
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(10.5)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    return table

def generate_thesis():
    """生成论文文档"""
    # 创建文档
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    
    # ===== 封面 =====
    # 学校名称
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('长沙学院')
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '黑体')
    run.font.size = Pt(26)
    run.font.bold = True
    p.paragraph_format.space_after = Pt(12)
    
    # 本科生毕业论文
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('本科生毕业论文')
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '黑体')
    run.font.size = Pt(22)
    run.font.bold = True
    p.paragraph_format.space_after = Pt(30)
    
    # 空行
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 论文标题
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('面向智能家居的多协议边缘智能网关设计')
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '黑体')
    run.font.size = Pt(18)
    run.font.bold = True
    p.paragraph_format.space_after = Pt(60)
    
    # 空行
    doc.add_paragraph()
    
    # 信息表格
    info_items = [
        ('学院', '计算机科学与工程学院'),
        ('专业', '物联网工程'),
        ('学生姓名', '黄志宇'),
        ('学号', 'B20220305131'),
        ('班级', '22物联01'),
        ('校内指导教师', '杨军'),
        ('职称', '高级工程师'),
    ]
    
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(f'{label}：{value}')
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '宋体')
        run.font.size = Pt(14)
        p.paragraph_format.line_spacing = 2.0
    
    # 日期
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('2026年5月')
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '宋体')
    run.font.size = Pt(14)
    
    # 分页
    doc.add_page_break()
    
    # ===== 摘要 =====
    add_heading_custom(doc, '摘  要', level=1)
    
    abstract_text = '''随着智能家居与物联网技术的迅猛发展，多协议设备的互联互通已成为制约用户体验提升的瓶颈，传统的单一协议或云端集中处理方案难以兼顾低时延、高兼容性与本地自治需求。本文旨在设计并实现一款面向智能家居场景的多协议边缘智能网关，以解决异构设备接入困难与协同效率低下的问题。本研究采用系统设计与实验验证相结合的方法，基于树莓派和ESP32等硬件平台，设计了一套支持Wi-Fi、BLE与MQTT协议的软件系统，并实现了协议适配、消息路由及本地联动等功能。通过搭建真实测试环境，对网关的功能完整性、传输时延及断网恢复能力进行了全面测试，结果表明系统可有效满足智能家居设备互联互通与边缘决策的核心需求。本研究期望能为边缘计算在消费级物联网中的落地提供低成本、高效率的参考方案，推动智能家居系统向更开放、更自主的方向发展。'''
    
    add_paragraph_custom(doc, abstract_text)
    
    # 关键词
    p = doc.add_paragraph()
    run = p.add_run('关键词：')
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '黑体')
    run.font.size = Pt(12)
    run.font.bold = True
    
    run = p.add_run('边缘智能网关；多协议互联；智能家居；MQTT；协议转换')
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    run.font.bold = True
    p.paragraph_format.line_spacing = 1.5
    
    # 分页
    doc.add_page_break()
    
    # ===== ABSTRACT =====
    add_heading_custom(doc, 'ABSTRACT', level=1)
    
    abstract_en = '''With the rapid development of smart home and Internet of Things technology, the interconnection of multi-protocol devices has become a bottleneck restricting the improvement of user experience. Traditional single-protocol or cloud-centric processing solutions are difficult to balance low latency, high compatibility, and local autonomy requirements. This paper aims to design and implement a multi-protocol edge intelligent gateway for smart home scenarios to solve the problems of heterogeneous device access difficulties and low collaboration efficiency. This research adopts a method combining system design and experimental verification, and designs a software system supporting Wi-Fi, BLE and MQTT protocols based on hardware platforms such as Raspberry Pi and ESP32, and implements protocol adaptation, message routing and local linkage functions. By building a real test environment, the functional integrity, transmission delay and network disconnection recovery capability of the gateway were comprehensively tested. The results show that the system can effectively meet the core needs of smart home device interconnection and edge decision-making. This study expects to provide a low-cost and high-efficiency reference scheme for the implementation of edge computing in consumer-grade IoT, and promote the development of smart home systems towards a more open and autonomous direction.'''
    
    add_paragraph_custom(doc, abstract_en, first_line_indent=False)
    
    # Keywords
    p = doc.add_paragraph()
    run = p.add_run('Keywords: ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    
    run = p.add_run('Edge Intelligent Gateway; Multi-protocol Interconnection; Smart Home; MQTT; Protocol Conversion')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    p.paragraph_format.line_spacing = 1.5
    
    # 分页
    doc.add_page_break()
    
    # ===== 目录（占位）=====
    add_heading_custom(doc, '目  录', level=1)
    
    toc_content = [
        '摘  要.......................................................... I',
        'ABSTRACT......................................................... II',
        '第1章 绪论....................................................... 1',
        '  1.1 研究背景与意义.............................................. 1',
        '    1.1.1 研究背景.............................................. 1',
        '    1.1.2 研究意义.............................................. 3',
        '  1.2 国内外研究概况.............................................. 3',
        '    1.2.1 国外研究现状.......................................... 3',
        '    1.2.2 国内研究现状.......................................... 4',
        '    1.2.3 研究述评.............................................. 5',
        '  1.3 研究内容.................................................... 6',
        '  1.4 研究方法.................................................... 6',
        '第2章 智能家居无线通信协议与关键技术分析.......................... 8',
        '  2.1 智能家居主流无线通信协议比较分析............................ 8',
        '  2.2 本系统支持协议与技术选型依据............................... 11',
        '  2.3 多协议共存与边缘侧协同可行性分析........................... 14',
        '第3章 多协议边缘网关总体方案与开发平台........................... 16',
        '  3.1 网关总体方案设计........................................... 16',
        '  3.2 硬件平台设计与实现......................................... 18',
        '  3.3 软件平台搭建与运行环境配置................................. 21',
        '第4章 多协议边缘网关软件系统设计与实现........................... 24',
        '  4.1 软件系统总体设计........................................... 24',
        '  4.2 统一数据传输模型与主题规范设计............................... 26',
        '  4.3 协议适配与数据收发模块设计实现............................. 28',
        '  4.4 协议转换与消息路由模块设计实现............................. 31',
        '  4.5 本地自治联动与离线补传机制实现............................. 33',
        '  4.6 设备管理与状态管理功能实现................................. 36',
        '第5章 多协议边缘网关系统测试与验证.............................. 39',
        '  5.1 测试环境与测试方案......................................... 39',
        '  5.2 系统功能测试与结果分析..................................... 41',
        '  5.3 系统性能测试与结果分析..................................... 43',
        '  5.4 系统稳定性测试与结果分析................................... 46',
        '  5.5 测试结果对比与达标分析..................................... 48',
        '第6章 总结与展望.................................................. 50',
        '  6.1 研究总结................................................. 50',
        '  6.2 研究不足................................................. 51',
        '  6.3 后续展望................................................. 52',
        '参考文献.......................................................... 54',
    ]
    
    for item in toc_content:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)
        p.paragraph_format.line_spacing = 1.5
    
    print("封面、摘要、目录添加完成")
    
    # 保存文档
    output_path = 'e:\\100_study\\120_Project\\CapstoneProject\\SmartHome_MultiProtocol_EdgeIntelligent_Gateway\\SH-MP-EG\\docs\\output\\面向智能家居的多协议边缘智能网关设计_论文_v2.docx'
    doc.save(output_path)
    print(f'文档已保存至: {output_path}')

if __name__ == '__main__':
    import docx
    from docx.oxml.ns import qn
    from docx.oxml import parse_xml
    generate_thesis()
