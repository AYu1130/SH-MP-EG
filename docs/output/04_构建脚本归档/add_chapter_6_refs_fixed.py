#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
添加第6章和参考文献（修复版）
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

def set_chinese_font(run, font_name='宋体', size=12, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold

def add_heading_custom(doc, text, level=1):
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    
    if level == 1:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(16)
        run.font.bold = True
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        heading.paragraph_format.space_before = Pt(24)
        heading.paragraph_format.space_after = Pt(18)
    elif level == 2:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(14)
        run.font.bold = True
        heading.paragraph_format.space_before = Pt(18)
        heading.paragraph_format.space_after = Pt(12)
    else:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(12)
        run.font.bold = True
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)
    
    heading.paragraph_format.line_spacing = 1.5
    return heading

def add_paragraph_custom(doc, text, first_line_indent=True, fontsize=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(fontsize)
    
    p.paragraph_format.line_spacing = 1.5
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(6)
    return p

# 打开文档
doc = Document('e:\\100_study\\120_Project\\CapstoneProject\\SmartHome_MultiProtocol_EdgeIntelligent_Gateway\\SH-MP-EG\\docs\\output\\论文_修复版_部分4.docx')

# ========== 第6章 ==========
doc.add_page_break()
add_heading_custom(doc, '第6章 总结与展望', level=1)

add_heading_custom(doc, '6.1 研究总结', level=2)

add_heading_custom(doc, '6.1.1 研究回顾与总括', level=3)

para_611 = [
    '本文针对智能家居场景中多协议设备互联互通这一核心问题，从系统工程的分析视角出发，设计并实现了一款面向智能家居的多协议边缘智能网关[1,6]。研究过程中，首先通过文献研究法梳理了国内外在智能家居网关、边缘计算、多协议融合等领域的研究现状，识别出现有方案在协议适配轻量化、本地自治能力、断网容错机制等方面的改进空间[7,11]。',
    
    '在此基础上，明确了本研究的核心目标：构建一套软硬件一体的边缘网关解决方案，实现Wi-Fi和BLE两种主流无线协议的统一接入，提供低于500毫秒的本地控制响应时延，支持网络中断场景下的离线自治和数据补传[8]。围绕这一目标，本文完成了硬件平台选型与搭建、软件系统架构设计、各功能模块实现、系统测试验证等全流程工作。',
    
    '通过原型系统的开发与实验验证，本文实现了预期的研究目标。硬件层面形成了以树莓派4B为主控、ESP32-S3和STM32F103为终端节点的异构架构；软件层面构建了包括协议适配、统一数据模型、消息路由、本地联动、离线缓存、Web管理等在内的完整功能体系[6,8]。系统测试结果表明，各项性能指标均达到或优于设计预期。'
]

for text in para_611:
    add_paragraph_custom(doc, text)

add_heading_custom(doc, '6.1.2 核心结论概括', level=3)

para_612 = [
    '通过本研究的系统设计与实验验证，形成以下核心结论：',
    
    '第一，本设计成功验证了基于树莓派和ESP32的异构多协议网关在功能上的完备性与可行性[5,6]。树莓派4B作为主控平台能够有效承载MQTT Broker、Node-RED、Flask Web服务等多种应用的并发运行；ESP32-S3和STM32F103作为终端节点提供了稳定的Wi-Fi和BLE接入能力。这一硬件组合在满足功能需求的同时保持了合理的成本控制。',
    
    '第二，本研究揭示了并发接入时协议转换引擎的性能瓶颈[6]。测试表明，当消息速率达到200条/秒以上时，系统开始出现丢包现象，主要由MQTT客户端库的消息队列处理能力限制导致。这一发现为未来优化工作指明了方向：可通过引入异步消息队列（如RabbitMQ）或优化MQTT客户端实现来提升高并发场景下的吞吐量。',
    
    '第三，本研究验证了采用统一数据模型和本地规则引擎实现边缘侧设备高效协同的有效路径[1,8]。通过定义标准化的JSON消息格式和层次化的MQTT主题规范，异构协议设备的数据实现了语义层面的互通；Node-RED可视化规则引擎降低了联动逻辑的开发和维护成本。'
]

for text in para_612:
    add_paragraph_custom(doc, text)

add_heading_custom(doc, '6.2 研究不足', level=2)

para_62 = [
    '客观审视本研究的局限性，主要包括以下几个方面：',
    
    '第一，受限于实验设备和测试条件，本研究未在大规模并发场景下进行极限压力测试[6]。当前测试环境最多同时接入5台Wi-Fi设备和3台BLE设备，与实际智能家居环境中可能出现的数十台设备同时接入的场景存在差距。在高密度设备部署场景下，射频干扰、网络拥塞、处理能力瓶颈等问题可能更为突出。',
    
    '第二，本研究仅对Wi-Fi和BLE两种协议进行了深度适配和验证[5,6]。虽然这两种协议覆盖了智能家居中大部分设备类型，但仍有相当数量的设备采用Zigbee、Z-Wave、Thread等协议，本系统尚不支持。协议支持的广度限制了系统在实际部署中的普适性。',
    
    '第三，系统的安全防护机制相对基础[11,12]。当前实现仅依赖Wi-Fi WPA2/3加密和MQTT基础认证，缺乏端到端的数据加密、设备身份认证、异常行为检测等高级安全功能。在智能家居数据隐私日益受到重视的背景下，安全防护能力有待加强。',
    
    '第四，Web管理台的用户界面设计较为简洁，功能相对基础[6]。设备发现、规则配置等操作仍需一定程度的技术背景，普通用户的上手体验有待优化。移动端适配尚未完成，限制了用户随时随地管理设备的可能性。'
]

for text in para_62:
    add_paragraph_custom(doc, text)

add_heading_custom(doc, '6.3 后续展望', level=2)

para_63 = [
    '基于本研究的成果与局限，未来工作可从以下方向展开：',
    
    '第一，在协议支持方面，后续研究可扩展对Zigbee、Z-Wave、Matter等协议的支持，实现真正意义上的全协议覆盖[1,4]。可探索在软件层面抽象出通用的协议驱动插件接口，允许第三方开发者贡献新协议的适配模块，构建开放的协议适配生态。',
    
    '第二，在智能化能力方面，可在网关中引入轻量级TinyML模型，通过对设备行为数据的边缘侧学习，实现从被动联动到主动智能预测的升级[7,8]。例如，通过学习用户的开关灯习惯，网关可主动预判用户行为并提前执行控制指令，进一步提升用户体验。',
    
    '第三，在安全防护方面，后续工作可探索基于TLS的端到端加密通信、基于数字证书的设备身份认证、基于规则的异常行为检测机制等[11,12]。还可研究边缘侧的数据脱敏和本地化处理策略，在满足服务需求的同时最大限度保护用户隐私。',
    
    '第四，在系统优化方面，可针对高并发场景进行性能优化，包括采用异步IO提升网络处理能力、引入消息队列削峰填谷、优化SQLite缓存机制等[6,8]。还可探索容器化部署方案，实现网关软件的快速部署和版本升级。',
    
    '综上所述，本研究为面向智能家居的多协议边缘智能网关设计提供了一个可行且有效的技术方案，验证了边缘计算在消费级物联网中的应用价值。随着技术的不断进步和需求的持续演进，边缘智能网关将在智能家居生态中扮演越来越重要的角色。'
]

for text in para_63:
    add_paragraph_custom(doc, text)

print("第6章完成")

# ========== 参考文献 ==========
doc.add_page_break()
add_heading_custom(doc, '参考文献', level=1)

# 参考文献列表（已移除16-18）
references = [
    '[1] 杨飞宇. 面向智能家居的边缘智能网关设计[D]. 华中科技大学, 2023.',
    '[2] Desai P, Sheth A, Anantharam P. Semantic gateway as a service architecture for IoT interoperability[C]//IEEE International Conference on Mobile Services. IEEE, 2015: 313-319.',
    '[3] 吴磊, 朱杰. 基于家庭多协议网关的通信协议设计[J]. 通信技术, 2017, 50(8): 1756-1762.',
    '[4] 王猛. 支持多协议的家庭网关的设计与实现[D]. 南京邮电大学, 2017.',
    '[5] 王琦锋. 基于MQTT和ESP-NOW的智能家居监测与联动系统[D]. 江南大学, 2023.',
    '[6] 刘亮亮, 田启川. 基于树莓派与MQTT的智能网关设计[J]. 物联网技术, 2024, 14(2): 73-77.',
    '[7] Stojkovic S, Jevtic B, Vukmirovic G, et al. Smart indoor gateway for home automation[C]//2018 26th Telecommunications Forum (TELFOR). IEEE, 2018: 1-4.',
    '[8] 王琦锋. 基于MQTT和ESP-NOW的智能家居监测与联动系统设计[D]. 江南大学, 2023.',
    '[9] Perera C, Zaslavsky A, Christen P, et al. Sensing-as-a-service model for smart cities supported by internet of things[J]. Transactions on Emerging Telecommunications Technologies, 2014, 25(1): 81-93.',
    '[10] Stojkovic S, Jevtic B, Vukmirovic G, et al. Indoor gateway for smart home system[C]//2018 26th Telecommunications Forum (TELFOR). IEEE, 2018: 420-423.',
    '[11] Li S, Da Xu L, Zhao S. The internet of things: a survey[J]. Information Systems Frontiers, 2015, 17(2): 243-259.',
    '[12] Li Q, Chen C, Xie Z, et al. SafeNet: A customizable malware detection system for smart home gateways[J]. IEEE Access, 2016, 4: 6345-6354.',
    '[13] Zhang L, Cao G, Wang L. Multi-protocol gateway design for smart home systems[J]. Journal of Network and Computer Applications, 2021, 174: 102983.',
    '[14] Singh M, Rajan M A, Shivraj V L, et al. Secure MQTT for internet of things (IoT)[C]//2015 Fifth International Conference on Communication Systems and Network Technologies. IEEE, 2015: 746-751.',
    '[15] Huh J H. Smart gateway for a social web of things (SWoT) integrated into a smart city system based on big data, context awareness, and the collective internet of things[J]. Sustainability, 2021, 13(3): 1321.',
]

for ref in references:
    p = doc.add_paragraph()
    run = p.add_run(ref)
    set_chinese_font(run, '宋体', 10.5)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.first_line_indent = Cm(-0.74)
    p.paragraph_format.space_after = Pt(3)

print("参考文献完成")

# ========== 致谢 ==========
doc.add_page_break()
add_heading_custom(doc, '致  谢', level=1)

para_thanks = [
    '时光荏苒，四年的本科学习即将画上句号。在这段求学旅程中，我收获了知识、友谊和成长，这一切离不开众多师长、同学和家人的支持与帮助。',
    
    '首先，衷心感谢我的校内指导教师杨军高级工程师。从选题、开题到论文撰写，杨老师始终给予悉心指导。杨老师严谨的治学态度、渊博的专业知识和耐心的教导让我受益匪浅。在遇到技术难题时，杨老师总能提供启发性的建议，帮助我找到解决方案。',
    
    '感谢计算机科学与工程学院的所有任课老师，是你们系统的专业课程培养了我的专业素养和工程实践能力。特别感谢物联网工程系的各位老师，在专业学习和科研训练中给予的指导和帮助。',
    
    '感谢实验室的同学们，在项目进行过程中与我相互交流、共同探讨，让我在团队协作中学会了沟通与分享。',
    
    '最后，感谢我的家人一直以来的理解和支持，是你们无私的爱让我能够专心学业，顺利完成学业。',
    
    '谨以此文献给所有关心和帮助过我的人！'
]

for text in para_thanks:
    add_paragraph_custom(doc, text)

print("致谢完成")

# 保存最终文档
doc.save('e:\\100_study\\120_Project\\CapstoneProject\\SmartHome_MultiProtocol_EdgeIntelligent_Gateway\\SH-MP-EG\\docs\\output\\面向智能家居的多协议边缘智能网关设计_修复版_完整版.docx')
print("最终文档保存完成！")
print("文件路径: docs\\output\\面向智能家居的多协议边缘智能网关设计_修复版_完整版.docx")
