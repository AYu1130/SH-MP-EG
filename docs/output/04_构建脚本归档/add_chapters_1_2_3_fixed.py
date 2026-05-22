#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
添加第1-3章内容（修复版）
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_chinese_font(run, font_name='宋体', size=12, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold

def add_heading_custom(doc, text, level=1):
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    
    if level == 1:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(16)
        run.font.bold = True
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        heading.paragraph_format.space_before = Pt(24)
        heading.paragraph_format.space_after = Pt(18)
    elif level == 2:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(14)
        run.font.bold = True
        heading.paragraph_format.space_before = Pt(18)
        heading.paragraph_format.space_after = Pt(12)
    else:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(12)
        run.font.bold = True
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)
    
    heading.paragraph_format.line_spacing = 1.5
    return heading

def add_paragraph_custom(doc, text, first_line_indent=True, fontsize=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(fontsize)
    
    p.paragraph_format.line_spacing = 1.5
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_figure_placeholder(doc, fig_num, caption, note=None):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(f'【在此处插入 图{fig_num}】')
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(11)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(f'图{fig_num} {caption}')
    set_chinese_font(run, '宋体', 10.5)
    
    if note:
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(f'({note})')
        set_chinese_font(run, '宋体', 9)
        p.paragraph_format.space_after = Pt(12)
    else:
        p.paragraph_format.space_after = Pt(12)

def add_table_caption(doc, caption):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(caption)
    set_chinese_font(run, '宋体', 10.5, True)
    p.paragraph_format.space_after = Pt(6)

def add_table_from_data(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(10)
                run.font.bold = True
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
            for p in row_cells[i].paragraphs:
                for run in p.runs:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(10)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    return table

# 打开文档
doc = Document('e:\\100_study\\120_Project\\CapstoneProject\\SmartHome_MultiProtocol_EdgeIntelligent_Gateway\\SH-MP-EG\\docs\\output\\论文_修复版_部分1.docx')

# ========== 第1章 绪论 ==========
doc.add_page_break()
add_heading_custom(doc, '第1章 绪论', level=1)

add_heading_custom(doc, '1.1 研究背景与意义', level=2)
add_heading_custom(doc, '1.1.1 研究背景', level=3)

para_111 = [
    '当前正处于物联网技术蓬勃发展的时代，人工智能技术与5G通信技术的深度融合正在重塑各个行业的面貌[1]。在"数字中国"战略的驱动下，智能家居领域正经历着从单品智能向全屋智能、从云端集中处理向边端协同计算的深刻转型。这一转型不仅体现在技术层面的革新，更反映了用户对家居体验智能化、个性化、实时化的内在诉求。智能家居系统需要整合照明控制、环境监测、安防预警、家电管理等多种功能，而这些功能往往依赖于不同厂商、不同协议的设备协同工作[2]。',
    
    '从产业应用的视角审视，智能家居市场呈现高度碎片化的特征。Wi-Fi、蓝牙低功耗（BLE）、Zigbee、Z-Wave等多种无线通信协议并存[3]，各大厂商围绕自身技术生态构建封闭的智能家居体系。用户在组建智能家居系统时，常常面临同一空间内不同品牌设备无法互联互通的困境，被迫购置多个专用网关，形成严重的"协议孤岛"现象。这种状况不仅增加了用户的经济负担，更使得跨品牌设备的协同控制成为不可能完成的任务[4]。',
    
    '在用户体验层面，现有智能家居系统普遍采用"终端-云端-终端"的通信模式，所有数据和控制指令都需要经过云端服务器中转[5]。这种模式在网络环境不稳定或完全中断的情况下会导致系统整体失效，且较高的传输时延难以满足实时性要求较高的场景需求。用户对于智能灯控、安防报警等场景，期望获得毫秒级的响应体验，而云端架构往往只能提供秒级的响应能力[6]。',
    
    '基于上述背景，部署于家庭网络边缘侧、能够同时接入多种协议设备并实现本地自治决策的多协议边缘智能网关，成为打破生态壁垒、解决时延与依赖问题的关键技术路径。通过在网关侧完成协议转换、数据融合与规则执行，不仅能够实现异构设备的无缝互联，更能在断网环境下维持核心功能的正常运转，显著提升系统的可靠性与用户体验[7]。'
]

for text in para_111:
    add_paragraph_custom(doc, text)

add_heading_custom(doc, '1.1.2 研究意义', level=3)

para_112 = [
    '本研究在理论层面具有一定的学术价值。现有文献中关于边缘计算网关的研究多聚焦于工业场景或城市级物联网应用[8]，针对消费级智能家居场景的轻量级边缘网关系统性设计方法讨论相对不足。本研究预期能为边缘计算网关的微服务架构设计、多协议数据模型的统一化理论、以及端-边-云协同计算的优化策略提供新的设计视角和实验证据。',
    
    '在工程实践层面，本研究形成的设计方案能够为嵌入式工程师、智能家居系统集成商提供一套具体、可复现的软硬件参考架构。该方案采用树莓派4B作为主控平台，结合ESP32-S3与STM32F103实现异构接入，整体硬件成本控制在合理范围内，具有良好的性价比优势。通过在网关侧实现协议转换与本地联动规则引擎，该方案能够直接降低智能家居系统的响应时延，测试表明本地控制响应时间可控制在500毫秒以内，MQTT端到端传输延迟不超过200毫秒[6]。在网络中断场景下，系统依靠本地规则引擎仍能维持智能联动功能的正常运行[7]。'
]

for text in para_112:
    add_paragraph_custom(doc, text)

add_heading_custom(doc, '1.2 国内外研究概况', level=2)
add_heading_custom(doc, '1.2.1 国外研究现状', level=3)

para_121 = [
    '国外在边缘计算与物联网关领域的研究起步较早，已形成较为成熟的理论体系和技术方案[7]。根据技术路线的差异，现有研究大致可分为两个主要方向。',
    
    '基于规范化中间件的多协议适配研究是国外学者关注的重点领域之一。Desai等人（2015）[7]提出的语义网关即服务架构，通过引入语义层抽象，实现了物联网设备间的互操作性。该架构在网关层部署语义解析引擎，将不同协议的数据映射为统一的资源描述框架，有效解决了异构设备的数据格式差异问题。Perera等人（2014）[9]则从传感即服务的角度，构建了支持多协议接入的智能城市物联网网关。',
    
    'Stojkovic等人（2018）[10]设计的智能家居网关采用了类似的分层思想，在网关内部实现Zigbee、Wi-Fi等协议的数据转换，支持异构网络节点间的双向透明传输。该研究通过制定统一的通信协议，设计了数据校验、确认重传等机制，保证了数据传输的正确性，网关的数据丢包率低于1.00%，通信时延低于115ms。',
    
    '基于边缘智能的自治与协同方法研究代表了另一重要方向。Ray（2018）[11]在物联网架构综述中系统分析了集中式云处理与边缘处理的优劣，指出边缘计算在降低传输时延、减少带宽占用、增强隐私保护方面具有显著优势。Li等人（2016）[12]从安全角度研究了边缘网关的漏洞检测机制，提出基于代码相似性分析的自动化检测方法。'
]

for text in para_121:
    add_paragraph_custom(doc, text)

add_heading_custom(doc, '1.2.2 国内研究现状', level=3)

para_122 = [
    '国内学者对于智能家居网关和多协议融合的研究已取得广泛成果，研究思路可大致分为两类[1]。',
    
    '在侧重云端平台与网关联动的研究中，杨飞宇（2023）[1]设计了面向智能家居的边缘智能网关，构建了"感知层-边缘层-网络层-应用层"四层架构，提出了基于模糊理论和层次分析法的二级数据融合算法。刘亮亮等人（2024）[6]基于树莓派与MQTT设计了智能网关，集成了服务代理、数据管理、协议转换等功能，实验表明在数千次传输中数据丢失率为零。',
    
    '王琦锋（2023）[4]基于MQTT和ESP-NOW设计了智能家居监测与联动系统，以ESP32为核心完成物联网节点开发，实现了设备本地联动和低功耗无线通信。该研究采用"云-边-端"三层架构，结合Docker技术和MQTT协议构建多协议接入方法[8]。',
    
    '在侧重轻量级协议转换硬件优化的研究中，吴磊和朱杰（2017）[3]设计了基于家庭多协议网关的通信协议，在协议数据单元中定义设备编号、协议类型、命令类型等字段，实现了对Wi-Fi、蓝牙、Zigbee等不同协议设备的统一控制。王猛（2017）[4]实现了支持Zigbee协议和蓝牙协议的多协议网关，采用C/S架构通过Wi-Fi和蓝牙模块实现对异构协议设备的集中远程控制。'
]

for text in para_122:
    add_paragraph_custom(doc, text)

add_heading_custom(doc, '1.2.3 研究述评', level=3)

para_123 = [
    '综合国内外研究成果可以发现，现有研究已经系统化地揭示了多协议网关的关键设计要素和技术实现路径，为本研究提供了宝贵的理论和实践基础[1,7,11]。然而，现有方案在若干方面仍存在改进空间。在研究内容上，现有工作多集中于特定协议组合的适配，对多协议动态切换和并发处理的机制设计关注较少，难以应对智能家居场景中设备类型多样化、接入时序不确定的复杂情况。',
    
    '此外，现有研究对网关在实际家庭网络环境下的长期稳定性、断网恢复能力等工程化指标的系统性验证相对不足[6]。因此，本文将从系统工程的视角，设计并实现一种结合Wi-Fi、BLE与MQTT的边缘智能网关，重点在协议适配模块的轻量化设计、统一数据模型的规范化定义、本地联动规则的灵活配置以及端到端性能的系统化验证等方面进行补充和改进[7]。'
]

for text in para_123:
    add_paragraph_custom(doc, text)

add_heading_custom(doc, '1.3 研究内容', level=2)

para_13 = [
    '本文以智能家居中多协议设备互联互通为核心问题，采用系统工程的分析视角与原型验证的研究方法[6]。通过对现有智能家居系统架构的深入分析，识别出协议异构、响应时延高、云端依赖强等关键痛点，明确边缘智能网关在解决这些问题中的核心作用。',
    
    '在此基础上，预期设计一套软硬件一体的多协议边缘网关解决方案。硬件层融合树莓派4B、ESP32-S3、STM32F103等多款处理器模块，形成异构物理接入能力：树莓派作为主控核心负责复杂的数据处理与规则执行，ESP32-S3承担Wi-Fi设备的接入与协议转换，STM32F103配合蓝牙模块实现BLE设备的低功耗接入[5,6]。软件层构建统一传输模型、协议适配引擎与本地联动规则等核心功能模块，实现异构协议数据的标准化封装、智能化路由与自动化处理。',
    
    '通过该设计，从理论和实践上论证边缘协同计算在提升智能家居系统兼容性、响应实时性及运行可靠性方面的价值，验证基于低成本嵌入式平台实现多协议边缘智能网关的技术可行性[1,8]。'
]

for text in para_13:
    add_paragraph_custom(doc, text)

add_heading_custom(doc, '1.4 研究方法', level=2)

add_heading_custom(doc, '(1) 文献研究法', level=3)
add_paragraph_custom(doc, '本文使用文献研究法，围绕多协议边缘网关的设计课题，广泛查阅国内外关于物联网协议、边缘计算架构及嵌入式系统设计的文献资料[1,7,11]。通过梳理主流的技术流派与实现方案，明晰当前多协议融合面临的技术瓶颈与前沿解决方案，为本研究的设计选型、架构规划及问题定义提供了坚实的理论框架。')

add_heading_custom(doc, '(2) 系统原型法', level=3)
add_paragraph_custom(doc, '本文采用系统工程中的原型开发模型，将整个网关系统分解为硬件平台构建、协议适配层开发、消息路由与联动应用设计等模块[8,6]。通过迭代式的开发与集成，最终构建出一个可运行的边缘网关原型系统。该方法旨在将理论框架转化为可验证的工程实体。')

add_heading_custom(doc, '(3) 实验验证法', level=3)
add_paragraph_custom(doc, '为评估所设计网关的实际效能，本文设计了包含功能、性能和稳定性测试在内的实验方案[2,6]。通过对系统传输时延、吞吐量、丢包率及断网恢复时间等关键指标的定量测量，并与现有方案或设计指标进行比对，客观地验证本系统设计方案的有效性与优越性。')

# ========== 第2章 ==========
doc.add_page_break()
add_heading_custom(doc, '第2章 智能家居无线通信协议与关键技术分析', level=1)

add_heading_custom(doc, '2.1 智能家居主流无线通信协议比较分析', level=2)

para_21 = [
    '智能家居领域的无线通信技术呈现出多元化发展的态势，不同协议在传输速率、功耗特性、网络拓扑、安全性及成本方面各具特色[3]。当前主流的无线通信协议主要包括Wi-Fi、蓝牙低功耗（BLE）、Zigbee、Z-Wave、Thread以及LoRaWAN等，这些协议分别适用于不同的应用场景和设备类型[4,6]。',
    
    'Wi-Fi协议以其高传输速率和广泛的设备兼容性成为智能家居中多媒体设备的首选通信方式[1,6]。基于IEEE 802.11系列标准，Wi-Fi在2.4GHz和5GHz频段提供高达数百Mbps的数据传输能力，能够满足视频监控、高清音频流等高带宽应用需求。然而，较高的功耗特性限制了其在电池供电传感器中的应用范围[5]。',
    
    '蓝牙低功耗（BLE）协议是经典蓝牙的演进版本，专注于低功耗、短距离的数据传输需求[4,10]。BLE采用跳频扩频技术，在2.4GHz频段通过40个信道实现抗干扰通信，其待机功耗可低至微安级别，非常适合温湿度传感器、门窗磁感应器等电池供电设备[3]。',
    
    'Zigbee协议基于IEEE 802.15.4标准，专为低功耗、低速率的传感器网络设计[1,4]。Zigbee支持星型、树型和网状网络拓扑，单个网络可容纳多达65000个节点，具有极强的网络扩展能力。其自组网、自愈合特性使得网络中的设备能够自动寻找最佳通信路径[7]。'
]

for text in para_21:
    add_paragraph_custom(doc, text)

# 表2-1
add_table_caption(doc, '表2-1 智能家居主流无线通信协议对比')

add_table_from_data(
    doc,
    ['对比维度', 'Wi-Fi', 'BLE', 'Zigbee', 'Z-Wave', 'LoRaWAN'],
    [
        ['传输速率', '11-600Mbps', '1-2Mbps', '20-250kbps', '9.6-100kbps', '0.3-50kbps'],
        ['功耗水平', '高', '极低', '低', '低', '极低'],
        ['传输距离', '50-100m', '10-100m', '10-100m', '30-100m', '2-15km'],
        ['网络拓扑', '星型', '星型/Mesh', '星型/树型/Mesh', '网状', '星型'],
        ['最大节点数', '受路由器限制', '无限制', '65000', '232', '约20万'],
        ['工作频段', '2.4/5GHz', '2.4GHz', '2.4GHz', 'Sub-1GHz', 'Sub-1GHz'],
        ['安全性', 'WPA3加密', 'AES-128', 'AES-128', 'AES-128', 'AES-128'],
        ['应用场景', '摄像头', '传感器', '照明', '安防', '远程传感器'],
        ['成本', '中', '低', '低', '中', '中'],
    ]
)

add_paragraph_custom(doc, '通过对比分析可以看出，不同协议在性能指标上存在显著差异[5]。Wi-Fi以高带宽见长但功耗较高，适合需要传输大量数据的设备；BLE以超低功耗著称，适合电池供电的小型传感器；Zigbee在网络容量和自组网能力方面优势明显[10]。这种多协议并存的局面决定了多协议网关的必要性[1,3]。')

# 添加系统架构图占位符
add_figure_placeholder(
    doc, 
    '2-1', 
    '系统整体架构图',
    '展示从终端设备经网关到应用层的完整数据流'
)

add_heading_custom(doc, '2.2 本系统支持协议与技术选型依据', level=2)

add_heading_custom(doc, 'Wi-Fi协议选型依据', level=3)
add_paragraph_custom(doc, 'Wi-Fi作为高带宽设备接入协议，主要服务于需要传输视频流、音频流或大批量数据的智能终端[1,6]。在本系统中，ESP32-S3模块通过Wi-Fi与网关进行通信，承担温湿度、光照等环境数据的采集与上报任务。选择Wi-Fi的原因在于：其一，Wi-Fi在家庭环境中普及率极高；其二，ESP32系列芯片集成了高性能Wi-Fi射频前端，在保持较低成本的同时提供稳定的无线连接能力[5,6]；其三，Wi-Fi的高速传输特性使得设备固件空中升级（OTA）成为可能[8]。')

add_heading_custom(doc, 'BLE协议选型依据', level=3)
add_paragraph_custom(doc, 'BLE作为低功耗传感器接入协议，主要服务于电池供电或采用能量采集方式工作的终端节点[3,4]。STM32F103配合BLE模块，实现蓝牙设备的扫描、连接与数据交互。选择BLE的原因在于：其一，BLE的超低功耗特性使其成为纽扣电池供电传感器的理想选择，单节CR2032电池可支持数年的持续工作[10]；其二，BLE的快速连接特性能够实现即时响应的场景需求；其三，BLE广播模式支持无连接的数据传输，简化了设备接入流程[7]。')

add_heading_custom(doc, 'MQTT协议选型依据', level=3)
add_paragraph_custom(doc, 'MQTT作为应用层统一消息总线，承担协议转换后的数据汇聚与分发职能[6,9]。MQTT是一种基于发布/订阅模式的轻量级消息传输协议，其设计初衷就是为低带宽、高延迟或不可靠网络环境下的物联网设备提供可靠的消息传输服务[1,12]。选择MQTT的原因在于：其一，MQTT的发布/订阅模式实现了数据生产者与消费者的解耦；其二，MQTT支持多级主题（Topic）层次结构，便于对海量设备进行逻辑分组和精细化管理[8]；其三，MQTT提供三种服务质量（QoS）等级；其四，MQTT的遗嘱消息（Last Will）和保留消息（Retained Message）机制能够有效处理设备异常离线场景[11]。')

add_heading_custom(doc, 'ETSI MEC架构的理论指导作用', level=3)
add_paragraph_custom(doc, '欧洲电信标准化协会（ETSI）提出的多接入边缘计算（MEC）参考架构为本系统设计提供了重要的理论框架[7]。MEC架构明确了边缘层在计算体系中的定位：作为网络边缘的分布式计算节点，提供贴近数据源的计算、存储和网络加速能力。')

add_paragraph_custom(doc, '参照ETSI MEC架构，本系统采用四层分层设计：感知层对应各类智能家居终端设备，负责环境数据采集和用户指令执行；边缘层即本文设计的边缘智能网关，承担协议适配、数据融合、规则执行等核心功能；网络层提供网关与外部网络（包括互联网和云平台）的连接能力；应用层面向最终用户提供可视化管理和控制界面[1,8]。这种分层架构实现了功能解耦和职责清晰划分[6]。')

add_heading_custom(doc, '2.3 多协议共存与边缘侧协同可行性分析', level=2)

add_heading_custom(doc, '物理层共存可行性分析', level=3)
add_paragraph_custom(doc, '本系统涉及的Wi-Fi和BLE协议均工作在2.4GHz ISM频段，频段范围为2400-2483.5MHz[5,10]。Wi-Fi在该频段划分14个信道（国内可用13个），每个信道带宽20MHz（或40MHz捆绑）；BLE划分40个信道，每个信道带宽2MHz。从频谱分配角度分析，两种协议的频谱资源存在重叠，理论上可能产生同频干扰[3]。')

add_paragraph_custom(doc, '然而，现代无线通信技术通过多种机制有效缓解了同频干扰问题[4]。Wi-Fi采用CSMA/CA（载波侦听多路访问/冲突避免）机制，在发送数据前检测信道状态，仅在信道空闲时启动传输；BLE采用自适应跳频技术，每次连接事件在37个数据信道间伪随机切换，有效分散干扰风险[10]。更为关键的是，本系统采用时分复用的方式协调两种协议的通信时序：网关主控制器通过调度算法为Wi-Fi和BLE模块分配不同的工作时段，避免两者在同一时刻进行射频收发操作[1,5]。')

# 添加时序图占位符
add_figure_placeholder(
    doc,
    '2-2',
    '跨协议联动时序图',
    '展示BLE光照传感器检测、网关处理、Wi-Fi灯具响应的完整时序流程'
)

add_heading_custom(doc, '上层协同调度可行性分析', level=3)
add_paragraph_custom(doc, '在应用层面，MQTT消息中间件为多协议设备的数据抽象和主题映射提供了坚实的技术基础[6,9,11]。无论设备通过Wi-Fi还是BLE接入网关，其上报的数据经过协议转换层处理后，均被封装为统一的JSON格式消息，并发布到规范化的MQTT主题上。')

add_paragraph_custom(doc, '本系统设计的主题命名遵循层次化原则，采用"home/区域/设备类型/设备ID/数据类型"的五级结构[8]。例如，位于客厅的温度传感器上报数据发布的主题可能为"home/livingroom/temperature/sensor_01/data"，而针对该设备的控制指令则发布到"home/livingroom/temperature/sensor_01/control"主题。这种层次化的主题设计使得设备管理具有清晰的逻辑结构[1,11]。')

add_paragraph_custom(doc, '跨协议联动场景的实现依赖于MQTT的发布/订阅机制[1,6,9]。以"光照自动控灯"场景为例：BLE光照传感器定期将环境亮度数据发布到指定主题；Node-RED规则引擎订阅该主题，当检测到亮度值低于设定阈值时，自动向Wi-Fi智能灯具的控制主题发布开启指令；灯具接收到指令后执行开灯动作。整个联动过程完全在本地网络中完成，无需云端参与，响应时延可控制在毫秒级别[8,7]。')

print("第1-2章完成")

# 保存
doc.save('e:\\100_study\\120_Project\\CapstoneProject\\SmartHome_MultiProtocol_EdgeIntelligent_Gateway\\SH-MP-EG\\docs\\output\\论文_修复版_部分2.docx')
print("部分2保存完成")
