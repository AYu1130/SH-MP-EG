#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
添加第3-4章内容（修复版）- 包含代码块
"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_chinese_font(run, font_name='宋体', size=12, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold

def add_heading_custom(doc, text, level=1):
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    
    if level == 1:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(16)
        run.font.bold = True
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        heading.paragraph_format.space_before = Pt(24)
        heading.paragraph_format.space_after = Pt(18)
    elif level == 2:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(14)
        run.font.bold = True
        heading.paragraph_format.space_before = Pt(18)
        heading.paragraph_format.space_after = Pt(12)
    else:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(12)
        run.font.bold = True
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)
    
    heading.paragraph_format.line_spacing = 1.5
    return heading

def add_paragraph_custom(doc, text, first_line_indent=True, fontsize=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(fontsize)
    
    p.paragraph_format.line_spacing = 1.5
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_figure_placeholder(doc, fig_num, caption, note=None):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(f'【在此处插入 图{fig_num}】')
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(11)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(f'图{fig_num} {caption}')
    set_chinese_font(run, '宋体', 10.5)
    
    if note:
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(f'({note})')
        set_chinese_font(run, '宋体', 9)
        p.paragraph_format.space_after = Pt(12)
    else:
        p.paragraph_format.space_after = Pt(12)

def add_table_caption(doc, caption):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(caption)
    set_chinese_font(run, '宋体', 10.5, True)
    p.paragraph_format.space_after = Pt(6)

def add_table_from_data(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(10)
                run.font.bold = True
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
            for p in row_cells[i].paragraphs:
                for run in p.runs:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(10)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    return table

def add_code_block(doc, code_lines, caption):
    """添加代码块 - 使用表格实现代码框效果"""
    # 代码标题
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(caption)
    set_chinese_font(run, '宋体', 10.5, True)
    p.paragraph_format.space_after = Pt(6)
    
    # 创建代码框表格
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False
    
    cell = table.rows[0].cells[0]
    cell.width = Inches(6)
    
    # 设置灰色背景
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F5F5F5')
    shading_elm.set(qn('w:val'), 'clear')
    cell._element.get_or_add_tcPr().append(shading_elm)
    
    # 清空默认内容
    cell.text = ''
    
    # 添加代码行
    for line in code_lines:
        p = cell.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Cm(0.2)
    
    doc.add_paragraph()

# 打开文档
doc = Document('e:\\100_study\\120_Project\\CapstoneProject\\SmartHome_MultiProtocol_EdgeIntelligent_Gateway\\SH-MP-EG\\docs\\output\\论文_修复版_部分2.docx')

# ========== 第3章 ==========
doc.add_page_break()
add_heading_custom(doc, '第3章 多协议边缘网关总体方案与开发平台', level=1)

add_heading_custom(doc, '3.1 网关总体方案设计', level=2)

para_31 = [
    '基于前文对无线通信协议特性和边缘计算架构的分析，本章提出系统的整体分层架构方案[1,8]。本系统采用四层分层架构，自上而下依次为感知层、边缘层、网络层和应用层，各层之间通过标准化接口进行数据交互，实现功能解耦和模块化设计[6]。',
    
    '感知层位于架构最底层，由各类智能家居终端设备组成，包括环境传感器（温湿度、光照、人体红外）、安防设备（门磁、烟雾报警器）、执行器（智能灯具、空调控制器）以及视频监控设备等[4,5]。这些设备通过Wi-Fi、BLE等协议与上层网关建立通信连接，负责采集环境状态数据并执行控制指令。',
    
    '边缘层是本文设计的核心，由边缘智能网关硬件及其上运行的软件系统构成[1,7]。网关内部进一步细分为协议适配模块、消息路由模块、规则引擎模块和本地存储模块。协议适配模块负责Wi-Fi和BLE协议的数据收发与协议转换；消息路由模块基于MQTT实现跨协议的消息转发；规则引擎模块执行本地联动逻辑；本地存储模块在断网时缓存待发送数据[6,8]。',
    
    '网络层提供网关与外部网络的连接能力，主要包括以太网和Wi-Fi两种接入方式[6]。在正常工作模式下，网关通过家庭路由器连接互联网，实现与云平台的远程通信；在网络中断场景下，网关切换至离线自治模式，依靠本地规则引擎维持核心功能的正常运转。',
    
    '应用层面向最终用户提供设备管理、数据可视化和远程控制等服务[1,9]。本系统开发了基于Python Flask的Web管理台，用户可通过浏览器访问网关的Web界面，查看设备在线状态、历史数据曲线、配置联动规则等。'
]

for text in para_31:
    add_paragraph_custom(doc, text)

# 添加模块结构图占位符
add_figure_placeholder(
    doc,
    '3-1',
    '网关软件模块层次结构图',
    '展示协议适配层、消息路由层、应用层的模块划分和调用关系'
)

add_heading_custom(doc, '3.2 硬件平台设计与实现', level=2)

add_heading_custom(doc, '3.2.1 主控硬件平台选型', level=3)
add_paragraph_custom(doc, '树莓派4B被选定为网关主控硬件平台，承担复杂的应用逻辑处理、MQTT Broker运行和Web服务提供等任务[6,8]。树莓派4B采用Broadcom BCM2711四核ARM Cortex-A72处理器，主频1.5GHz，配备4GB LPDDR4内存，提供千兆以太网接口、802.11ac Wi-Fi和蓝牙5.0支持。选择树莓派4B的原因在于：其一，成熟的Linux生态系统（Raspberry Pi OS）便于部署EMQX、Node-RED等开源软件；其二，丰富的GPIO接口便于连接外设；其三，性价比较高[5]。')

add_heading_custom(doc, '3.2.2 Wi-Fi接入模块设计', level=3)
add_paragraph_custom(doc, 'ESP32-S3负责Wi-Fi设备的接入与数据处理，作为感知层的Wi-Fi终端模拟器和数据采集节点[5,6]。ESP32-S3是乐鑫科技推出的高性能Wi-Fi+蓝牙SoC，集成Xtensa LX7双核处理器，支持802.11 b/g/n Wi-Fi协议。在本系统中，ESP32-S3通过家庭Wi-Fi网络与树莓派网关建立TCP连接，模拟温湿度传感器的数据上报行为。')

# ESP32 WiFi代码
wifi_code = [
    'import network',
    'import socket',
    'import time',
    'import json',
    'from machine import Pin, ADC',
    '',
    'class WiFiSensorNode:',
    '    """Wi-Fi传感器节点类 - ESP32-S3实现"""',
    '    ',
    '    def __init__(self, ssid, password, gateway_ip, port):',
    '        self.ssid = ssid',
    '        self.password = password',
    '        self.gateway_ip = gateway_ip',
    '        self.port = port',
    '        self.wlan = network.WLAN(network.STA_IF)',
    '        self.sock = None',
    '        self.device_id = "wifi_temp_001"',
    '        self.reconnect_count = 0',
    '    ',
    '    def connect_wifi(self):',
    '        """连接Wi-Fi网络"""',
    '        self.wlan.active(True)',
    '        self.wlan.connect(self.ssid, self.password)',
    '        while not self.wlan.isconnected():',
    '            time.sleep(1)',
    '        print(f"WiFi connected: {self.wlan.ifconfig()}")',
    '    ',
    '    def connect_gateway(self):',
    '        """建立与网关的TCP连接"""',
    '        self.sock = socket.socket()',
    '        self.sock.connect((self.gateway_ip, self.port))',
    '        # 发送设备注册信息',
    '        reg_msg = json.dumps({',
    '            "device_id": self.device_id,',
    '            "type": "temperature_sensor",',
    '            "protocol": "wifi"',
    '        })',
    '        self.sock.send(reg_msg.encode())',
    '    ',
    '    def read_sensor(self):',
    '        """读取温度传感器数据（模拟）"""',
    '        adc = ADC(Pin(34))',
    '        adc.atten(ADC.ATTN_11DB)',
    '        raw = adc.read()',
    '        temp = 20.0 + (raw / 4095.0) * 15.0',
    '        return round(temp, 2)',
    '    ',
    '    def send_data(self, temp):',
    '        """发送数据到网关"""',
    '        msg = {',
    '            "device_id": self.device_id,',
    '            "timestamp": int(time.time()),',
    '            "type": "temperature",',
    '            "value": temp,',
    '            "unit": "Celsius"',
    '        }',
    '        self.sock.send(json.dumps(msg).encode())',
    '        print(f"Sent: {msg}")',
    '    ',
    '    def run(self):',
    '        self.connect_wifi()',
    '        self.connect_gateway()',
    '        while True:',
    '            try:',
    '                temp = self.read_sensor()',
    '                self.send_data(temp)',
    '                time.sleep(5)',
    '            except Exception as e:',
    '                print(f"Error: {e}")',
    '                self.reconnect_count += 1',
    '                time.sleep(10)',
    '                self.connect_gateway()',
]

add_code_block(doc, wifi_code, '代码清单3-1 ESP32-S3 Wi-Fi传感器节点核心代码')

add_heading_custom(doc, '3.2.3 BLE接入模块设计', level=3)
add_paragraph_custom(doc, 'STM32F103配合BLE模块负责蓝牙低功耗设备的接入[5]。STM32F103是意法半导体推出的主流ARM Cortex-M3微控制器，主频72MHz，具有低功耗、高性价比的特点。在本系统中，STM32F103通过UART与蓝牙模块通信，实现BLE设备的扫描、连接和数据交互功能。')

# 添加BLE代码占位符说明
add_paragraph_custom(doc, 'BLE模块的固件实现包括GATT客户端初始化、服务发现、特征值读写等核心功能。由于篇幅限制，具体代码实现详见附录或项目源码仓库。')

add_heading_custom(doc, '3.3 软件平台搭建与运行环境配置', level=2)

add_heading_custom(doc, '3.3.1 操作系统选择与配置', level=3)
add_paragraph_custom(doc, '树莓派4B运行Raspberry Pi OS（基于Debian的Linux发行版）作为网关的基础操作系统[6,8]。Raspberry Pi OS针对树莓派硬件进行了深度优化，提供完善的硬件驱动支持和丰富的软件包生态。系统安装完成后，需要进行以下配置：启用SSH远程访问以便维护；配置静态IP地址确保网关地址稳定；安装Python 3.9及以上版本作为开发环境基础。')

add_heading_custom(doc, '3.3.2 EMQX消息中间件部署', level=3)
add_paragraph_custom(doc, 'EMQX是一款高性能的开源MQTT消息代理服务器，支持百万级并发连接[11]。本系统采用EMQX 5.x版本作为网关内部的MQTT Broker。部署步骤包括：添加EMQX官方软件源；执行安装命令；修改配置文件设置监听端口和认证方式；启动服务并设置开机自启。配置完成后，EMQX默认在1883端口提供MQTT服务，在18083端口提供Web管理界面[6,8]。')

add_heading_custom(doc, '3.3.3 Node-RED规则引擎配置', level=3)
add_paragraph_custom(doc, 'Node-RED是基于Node.js的可视化编程工具，适合快速构建物联网数据流和自动化规则[8]。安装Node-RED后，通过浏览器访问1880端口即可进入可视化编辑界面。在本系统中，Node-RED通过mqtt in节点订阅传感器数据主题，通过function节点编写联动逻辑，通过mqtt out节点下发控制指令。')

# 添加Node-RED流程图占位符
add_figure_placeholder(
    doc,
    '3-2',
    'Node-RED联动规则流程图',
    '展示光照-灯具联动规则的节点连接配置'
)

print("第3章完成")

# ========== 第4章 ==========
doc.add_page_break()
add_heading_custom(doc, '第4章 多协议边缘网关软件系统设计与实现', level=1)

add_heading_custom(doc, '4.1 软件系统总体设计', level=2)

para_41 = [
    '本章详细阐述网关软件系统的模块化设计思路与实现细节[6,8]。软件系统采用模块化架构设计，将复杂的网关功能分解为若干相对独立、职责单一的模块，便于独立开发、测试和维护。',
    
    '从功能划分角度，软件系统包括五大核心模块：协议适配与数据收发模块、统一数据传输模型模块、协议转换与消息路由模块、本地自治联动模块、设备管理与状态管理模块[1,6]。协议适配与数据收发模块负责与Wi-Fi、BLE终端建立物理连接，实现原始数据的收发；统一数据传输模型模块定义标准化的JSON消息格式；协议转换与消息路由模块实现异构协议数据向统一格式的转换和MQTT主题的路由；本地自治联动模块基于Node-RED和SQLite实现规则编排和离线缓存；设备管理与状态管理模块提供Web管理界面[8]。'
]

for text in para_41:
    add_paragraph_custom(doc, text)

# 添加软件架构图占位符
add_figure_placeholder(
    doc,
    '4-1',
    '软件系统模块调用关系图',
    '展示各模块间的数据流向和调用关系'
)

add_heading_custom(doc, '4.2 统一数据传输模型与主题规范设计', level=2)

add_heading_custom(doc, '4.2.1 统一JSON数据格式定义', level=3)
add_paragraph_custom(doc, '为实现异构协议设备的数据互通，本系统定义了一套统一的JSON数据传输格式[1,6]。该格式包含以下标准字段：device_id（设备唯一标识符）、timestamp（数据生成时间戳）、data_type（数据类型）、payload（业务数据载荷）、protocol（原始协议类型）、qos（服务质量等级）。')

# JSON数据格式示例
json_code = [
    '{',
    '    "device_id": "wifi_temp_001",',
    '    "timestamp": 1701234567,',
    '    "data_type": "temperature",',
    '    "protocol": "wifi",',
    '    "qos": 1,',
    '    "payload": {',
    '        "value": 24.5,',
    '        "unit": "Celsius",',
    '        "location": "livingroom"',
    '    }',
    '}'
]

add_code_block(doc, json_code, '代码清单4-1 统一JSON数据格式示例')

add_heading_custom(doc, '4.2.2 MQTT主题命名规范', level=3)
add_paragraph_custom(doc, '本系统采用层次化的MQTT主题命名规范，遵循"home/区域/设备类型/设备ID/数据类型"的五级结构[6,11]。区域字段标识设备所在物理空间，如bedroom、livingroom、kitchen等；设备类型字段标识设备类别，如temperature、light、door等；设备ID字段是设备的唯一标识；数据类型字段区分数据上报（data）和控制指令（control）[8]。')

# 主题规范表格
add_table_caption(doc, '表4-1 MQTT主题命名规范示例')
add_table_from_data(
    doc,
    ['主题层级', '含义', '示例值'],
    [
        ['第1层 home', '根命名空间', 'home'],
        ['第2层 区域', '设备所在房间', 'livingroom, bedroom'],
        ['第3层 设备类型', '设备功能类别', 'temperature, light'],
        ['第4层 设备ID', '设备唯一标识', 'sensor_001, lamp_002'],
        ['第5层 数据类型', '数据方向', 'data, control, status'],
    ]
)

add_heading_custom(doc, '4.3 协议适配与数据收发模块设计实现', level=2)

add_heading_custom(doc, '4.3.1 Wi-Fi适配模块实现', level=3)
add_paragraph_custom(doc, 'Wi-Fi适配模块运行在树莓派网关侧，作为TCP服务器监听来自ESP32终端的连接请求[6]。模块采用多线程架构，主线程负责监听端口，每当有新连接建立时，创建独立的工作线程处理该设备的通信。')

# WiFiAdapter核心代码
wifi_adapter_code = [
    'import socket',
    'import threading',
    'import json',
    'import logging',
    'from datetime import datetime',
    '',
    'logger = logging.getLogger(__name__)',
    '',
    'class WiFiAdapter:',
    '    """Wi-Fi协议适配器 - TCP服务端实现"""',
    '    ',
    '    def __init__(self, host="0.0.0.0", port=8888):',
    '        self.host = host',
    '        self.port = port',
    '        self.devices = {}',
    '        self.server_socket = socket.socket(',
    '            socket.AF_INET, socket.SOCK_STREAM)',
    '        self.server_socket.setsockopt(',
    '            socket.SOL_SOCKET, socket.SO_REUSEADDR, 1)',
    '        self.on_message_callback = None',
    '        self._running = False',
    '    ',
    '    def set_message_callback(self, callback):',
    '        """设置消息回调函数"""',
    '        self.on_message_callback = callback',
    '    ',
    '    def start(self):',
    '        """启动TCP服务器"""',
    '        self.server_socket.bind((self.host, self.port))',
    '        self.server_socket.listen(10)',
    '        self._running = True',
    '        threading.Thread(target=self._accept_loop, daemon=True).start()',
    '        logger.info(f"WiFi server started on {self.host}:{self.port}")',
    '    ',
    '    def _accept_loop(self):',
    '        """接受连接的主循环"""',
    '        while self._running:',
    '            try:',
    '                client_socket, address = self.server_socket.accept()',
    '                logger.info(f"New connection from {address}")',
    '                device_thread = threading.Thread(',
    '                    target=self._handle_device,',
    '                    args=(client_socket, address),',
    '                    daemon=True)',
    '                device_thread.start()',
    '            except Exception as e:',
    '                logger.error(f"Accept error: {e}")',
    '    ',
    '    def _handle_device(self, client_socket, address):',
    '        """处理单个设备的通信"""',
    '        device_id = None',
    '        try:',
    '            # 首次通信进行设备认证',
    '            data = client_socket.recv(1024)',
    '            reg_info = json.loads(data.decode())',
    '            device_id = reg_info.get("device_id")',
    '            self.devices[device_id] = client_socket',
    '            logger.info(f"Device {device_id} registered")',
    '            ',
    '            # 持续接收数据',
    '            while self._running:',
    '                data = client_socket.recv(1024)',
    '                if not data:',
    '                    break',
    '                self._process_data(device_id, data)',
    '        except Exception as e:',
    '            logger.error(f"Device handler error: {e}")',
    '        finally:',
    '            if device_id:',
    '                self.devices.pop(device_id, None)',
    '            client_socket.close()',
    '    ',
    '    def _process_data(self, device_id, data):',
    '        """处理接收到的数据"""',
    '        try:',
    '            raw_data = json.loads(data.decode())',
    '            unified_msg = {',
    '                "device_id": device_id,',
    '                "timestamp": raw_data.get("timestamp", 0),',
    '                "data_type": raw_data.get("type", "unknown"),',
    '                "protocol": "wifi",',
    '                "payload": raw_data',
    '            }',
    '            if self.on_message_callback:',
    '                self.on_message_callback(unified_msg)',
    '        except json.JSONDecodeError as e:',
    '            logger.error(f"Invalid JSON from {device_id}: {e}")',
]

add_code_block(doc, wifi_adapter_code, '代码清单4-2 Wi-Fi适配器核心实现代码')

add_heading_custom(doc, '4.3.2 BLE适配模块实现', level=3)
add_paragraph_custom(doc, 'BLE适配模块采用Python Bleak库实现BLE设备的异步扫描、连接和通知订阅[10]。Bleak是跨平台的BLE客户端库，支持Windows、Linux和macOS，在本系统中运行于Raspberry Pi OS之上。')

# BLE适配器代码
ble_code = [
    'import asyncio',
    'from bleak import BleakScanner, BleakClient',
    'import json',
    'import logging',
    '',
    'logger = logging.getLogger(__name__)',
    'UART_TX_CHAR_UUID = "6e400003-b5a3-f393-e0a9-e50e24dcca9e"',
    '',
    'class BLEAdapter:',
    '    """BLE协议适配器 - GATT客户端实现"""',
    '    ',
    '    def __init__(self, gateway):',
    '        self.gateway = gateway',
    '        self.devices = {}',
    '        self.connected_devices = {}',
    '        self.scanner = BleakScanner()',
    '        self._scanning = False',
    '    ',
    '    async def start_scan(self):',
    '        """启动BLE设备扫描"""',
    '        self.scanner.register_detection_callback(',
    '            self._on_device_detected)',
    '        await self.scanner.start()',
    '        self._scanning = True',
    '        logger.info("BLE scanning started")',
    '    ',
    '    def _on_device_detected(self, device, adv_data):',
    '        """扫描到设备时的回调"""',
    '        device_id = self._extract_device_id(adv_data)',
    '        if device_id and device_id.startswith("ble_"):',
    '            self.devices[device_id] = {',
    '                "address": device.address,',
    '                "rssi": adv_data.rssi',
    '            }',
    '            logger.info(f"Detected BLE device: {device_id}")',
    '            # 异步连接设备',
    '            asyncio.create_task(',
    '                self._connect_device(device.address, device_id))',
    '    ',
    '    async def _connect_device(self, address, device_id):',
    '        """连接指定BLE设备"""',
    '        try:',
    '            client = BleakClient(address)',
    '            await client.connect()',
    '            self.connected_devices[device_id] = client',
    '            logger.info(f"Connected to {device_id}")',
    '            ',
    '            # 订阅通知',
    '            await client.start_notify(',
    '                UART_TX_CHAR_UUID,',
    '                lambda s, d: self._on_notification(device_id, s, d))',
    '        except Exception as e:',
    '            logger.error(f"BLE connect error for {device_id}: {e}")',
    '    ',
    '    def _on_notification(self, device_id, sender, data):',
    '        """收到BLE通知时的回调"""',
    '        try:',
    '            payload = self._parse_ble_data(data)',
    '            unified_msg = {',
    '                "device_id": device_id,',
    '                "timestamp": int(time.time()),',
    '                "data_type": payload.get("type", "unknown"),',
    '                "protocol": "ble",',
    '                "payload": payload',
    '            }',
    '            self.gateway.on_device_message(unified_msg)',
    '        except Exception as e:',
    '            logger.error(f"Notification parse error: {e}")',
]

add_code_block(doc, ble_code, '代码清单4-3 BLE适配器核心实现代码')

# 添加BLE流程图占位符
add_figure_placeholder(
    doc,
    '4-2',
    'BLE适配模块程序流程图',
    '展示扫描、连接、通知订阅、数据处理的完整流程'
)

print("第4章前半部分完成")

# 保存
doc.save('e:\\100_study\\120_Project\\CapstoneProject\\SmartHome_MultiProtocol_EdgeIntelligent_Gateway\\SH-MP-EG\\docs\\output\\论文_修复版_部分3.docx')
print("部分3保存完成")
