#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
添加第4-6章内容
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

def add_heading_custom(doc, text, level=1):
    """添加自定义标题"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    if level == 1:
        run.font.size = Pt(16)
        run.font.bold = True
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        heading.paragraph_format.space_before = Pt(24)
        heading.paragraph_format.space_after = Pt(18)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.bold = True
        heading.paragraph_format.space_before = Pt(18)
        heading.paragraph_format.space_after = Pt(12)
    else:
        run.font.size = Pt(12)
        run.font.bold = True
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)
    
    heading.paragraph_format.line_spacing = 1.5
    return heading

def add_paragraph_custom(doc, text, first_line_indent=True):
    """添加自定义段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    
    p.paragraph_format.line_spacing = 1.5
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(6)
    
    return p

def add_table_from_data(doc, headers, rows):
    """添加表格"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # 表头
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(10)
                run.font.bold = True
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 数据行
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
            for p in row_cells[i].paragraphs:
                for run in p.runs:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(10)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    return table

def main():
    # 打开现有文档
    doc = Document('e:\\100_study\\120_Project\\CapstoneProject\\SmartHome_MultiProtocol_EdgeIntelligent_Gateway\\SH-MP-EG\\docs\\output\\面向智能家居的多协议边缘智能网关设计_论文_v2.docx')
    
    # ========== 第4章 ==========
    doc.add_page_break()
    add_heading_custom(doc, '第4章 多协议边缘网关软件系统设计与实现', level=1)
    
    # 4.1
    add_heading_custom(doc, '4.1 软件系统总体设计', level=2)
    add_heading_custom(doc, '模块化设计思想', level=3)
    
    texts_41 = [
        '本系统软件设计遵循模块化、分层化的架构原则，将复杂的多协议网关功能分解为若干高内聚、低耦合的功能模块。每个模块承担明确的职责，通过定义良好的接口进行交互，便于独立开发、测试和维护。',
        '软件架构自上而下划分为四个层次：应用服务层、业务逻辑层、协议适配层和硬件抽象层。应用服务层面向用户提供交互接口，包括Web管理后台和RESTful API；业务逻辑层实现设备管理、联动规则、数据缓存等核心功能；协议适配层处理异构协议的解析与转换；硬件抽象层封装底层硬件接口，向上层提供统一的操作接口。'
    ]
    
    for text in texts_41:
        add_paragraph_custom(doc, text)
    
    # 表4-1
    p = doc.add_paragraph()
    run = p.add_run('表4-1 软件模块层次结构')
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(10.5)
    run.font.bold = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    headers = ['层次', '模块名称', '功能描述', '技术实现']
    rows = [
        ['应用服务层', 'Web管理台', '提供可视化设备管理界面', 'Flask + HTML/JavaScript'],
        ['应用服务层', 'REST API', '提供外部系统对接接口', 'Flask-RESTful'],
        ['业务逻辑层', '设备管理器', '设备注册、发现、状态维护', 'Python Class'],
        ['业务逻辑层', '联动引擎', '本地自动化规则执行', 'Node-RED'],
        ['业务逻辑层', '数据缓存', '离线数据存储与补传', 'SQLite'],
        ['协议适配层', 'Wi-Fi适配器', 'TCP/UDP Socket数据收发', 'Python Socket'],
        ['协议适配层', 'BLE适配器', '蓝牙GATT通信管理', 'Bleak库'],
        ['协议适配层', 'MQTT适配器', 'MQTT消息发布/订阅', 'Paho-MQTT'],
        ['协议适配层', '协议转换器', '数据格式转换与封装', 'JSON Schema'],
        ['硬件抽象层', 'GPIO控制', '树莓派GPIO接口操作', 'RPi.GPIO'],
        ['硬件抽象层', '串口通信', 'UART设备通信', 'PySerial'],
    ]
    add_table_from_data(doc, headers, rows)
    
    add_paragraph_custom(doc, '模块间的依赖关系遵循单向原则：上层模块可调用下层模块的接口，但下层模块不依赖上层模块的具体实现。这种设计保证了核心功能（如协议适配）的稳定性，同时允许上层应用（如Web界面）独立演进。')
    
    # 4.2
    add_heading_custom(doc, '4.2 统一数据传输模型与主题规范设计', level=2)
    add_heading_custom(doc, '通用JSON数据格式定义', level=3)
    
    add_paragraph_custom(doc, '为实现异构协议数据的标准化表示，本系统定义了一套通用的JSON数据格式。所有设备上报的数据和网关转发的消息均采用此格式封装，确保数据处理逻辑的一致性和可扩展性。')
    
    # 表4-2
    p = doc.add_paragraph()
    run = p.add_run('表4-2 通用JSON数据格式字段定义')
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(10.5)
    run.font.bold = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    headers = ['字段名', '数据类型', '必填', '说明']
    rows = [
        ['device_id', 'String', '是', '设备唯一标识符'],
        ['device_type', 'String', '是', '设备类型'],
        ['protocol', 'String', '是', '接入协议'],
        ['timestamp', 'Integer', '是', '数据采集时间戳'],
        ['location', 'String', '否', '设备位置信息'],
        ['data', 'Object', '是', '传感器数据对象'],
        ['status', 'String', '是', '设备状态'],
        ['battery', 'Integer', '否', '电池电量百分比'],
    ]
    add_table_from_data(doc, headers, rows)
    
    # 4.3-4.6 简要内容
    add_heading_custom(doc, '4.3 协议适配与数据收发模块设计实现', level=2)
    add_paragraph_custom(doc, 'Wi-Fi终端接入模块负责管理ESP32-S3设备通过Wi-Fi网络与网关的通信连接。模块采用TCP Socket Server模式运行，监听指定端口（默认8888），等待设备连接建立。模块由主监听线程和多个设备处理线程组成，主监听线程持续监听端口，当检测到新的设备连接请求时，验证设备身份，身份验证通过后创建独立的设备处理线程。')
    
    add_heading_custom(doc, '4.4 协议转换与消息路由模块设计实现', level=2)
    add_paragraph_custom(doc, '协议转换模块的核心功能是将Wi-Fi和BLE协议接收的原始数据，按照JSON规范进行标准化封装。转换过程包括字段映射、数据类型转换、时间戳统一和单位标准化四个步骤。字段映射规则根据设备类型的不同有所差异，协议转换器维护一个设备类型到字段映射规则的查找表，根据消息中的device_type字段选择对应的映射规则。')
    
    add_heading_custom(doc, '4.5 本地自治联动与离线补传机制实现', level=2)
    add_paragraph_custom(doc, 'Node-RED作为本地规则引擎，通过可视化节点编排实现设备联动的灵活配置。系统预置多种常用节点，用户通过拖拽连接即可构建自动化流程。为实现断网自治和数据完整性保障，系统引入SQLite本地数据库，在网络中断时缓存无法上传的数据，待网络恢复后按时间戳顺序进行补传。')
    
    add_heading_custom(doc, '4.6 设备管理与状态管理功能实现', level=2)
    add_paragraph_custom(doc, 'Web管理后台基于Flask框架构建，提供RESTful API接口和HTML前端页面。采用应用工厂模式组织代码，便于测试和配置管理。设备在线状态通过心跳超时机制进行判断，每个设备需定期发送心跳消息或数据上报消息，Web后端维护最后活跃时间戳，当超过设定阈值未收到消息时，将设备状态标记为离线。')
    
    # ========== 第5章 ==========
    doc.add_page_break()
    add_heading_custom(doc, '第5章 多协议边缘网关系统测试与验证', level=1)
    
    # 5.1
    add_heading_custom(doc, '5.1 测试环境与测试方案', level=2)
    add_paragraph_custom(doc, '测试硬件设备包括：树莓派4B（4GB RAM）作为网关主控，ESP32-S3开发板作为Wi-Fi终端节点，STM32F103+HC-08模块作为BLE终端节点，SHT30温湿度模块和BH1750光照模块作为传感器，继电器模块作为执行器。软件环境包括Raspberry Pi OS、Python 3.9、EMQX 5.0、Node-RED 3.0。测试工具包括MQTT.fx、Wireshark、mosquitto_pub/sub、iperf3等。')
    
    # 5.2
    add_heading_custom(doc, '5.2 系统功能测试与结果分析', level=2)
    add_paragraph_custom(doc, '功能测试覆盖了设备接入、数据上报、跨协议联动、Web管理、离线缓存等核心功能。设备接入测试验证了Wi-Fi和BLE两种协议的设备能够正常发现、连接并注册到网关。数据上报测试持续运行30分钟，数据解析成功率100%。跨协议联动测试中，光照自动控灯场景响应时间约350毫秒，验证了跨协议联动的端到端功能。离线缓存测试模拟网络中断场景，缓存的15条消息在3秒内全部成功补传。')
    
    # 5.3
    add_heading_custom(doc, '5.3 系统性能测试与结果分析', level=2)
    add_paragraph_custom(doc, '端到端传输时延测试结果显示，Wi-Fi传输时延平均值45.2ms，BLE时延平均值78.6ms，均满足设计指标要求。跨节点联动响应时延平均136.2ms，满足500ms的设计指标。并发接入能力测试表明，系统在设计指标要求的2台以上设备并发之上仍有较大余量，测试至10个并发设备时系统仍能稳定运行。丢包率测试显示Wi-Fi丢包率0.06%，BLE丢包率0.14%，数据准确率接近100%。')
    
    # 5.4
    add_heading_custom(doc, '5.4 系统稳定性测试与结果分析', level=2)
    add_paragraph_custom(doc, '7×24小时长时运行测试中，系统保持稳定运行，无崩溃或重启。CPU负载稳定在12%左右，峰值负载不超过35%。内存占用呈现缓慢增长趋势，7天累计增长17MB，日均增长约2.4MB。断网恢复能力测试显示，断网期间本地联动功能完全不受影响，SQLite缓存机制工作正常，所有测试场景的补传成功率均达到100%，30分钟断网测试中180条缓存数据的补传在23秒内完成。')
    
    # 5.5
    add_heading_custom(doc, '5.5 测试结果对比与达标分析', level=2)
    add_paragraph_custom(doc, '所有设计指标均达到或超过预期要求。本地控制响应时间136.2ms（指标≤500ms），MQTT端到端传输延迟45.2ms（指标≤200ms），同时稳定接入终端数10个（指标≥2个），数据准确率99.86%（指标≥95%）。与主流竞品方案相比，本设计方案的优势在于完全开源的架构和高度的可定制性，断网自治能力经过完整测试验证，在无外网环境下仍可维持全部本地功能。')
    
    # ========== 第6章 ==========
    doc.add_page_break()
    add_heading_custom(doc, '第6章 总结与展望', level=1)
    
    # 6.1
    add_heading_custom(doc, '6.1 研究总结', level=2)
    add_heading_custom(doc, '(1) 研究回顾与总括', level=3)
    add_paragraph_custom(doc, '本文针对智能家居领域多协议设备互联互通困难、云端架构响应时延高、断网可用性差等核心问题，设计并实现了一款面向智能家居场景的多协议边缘智能网关。研究采用系统工程的分析视角与原型验证相结合的方法，从协议分析、架构设计、软硬件实现到系统测试，构建了完整的技术方案体系。')
    
    add_heading_custom(doc, '(2) 核心结论概括', level=3)
    add_paragraph_custom(doc, '本设计成功验证了基于树莓派和ESP32的异构多协议网关在功能上的完备性与可行性。系统同时接入Wi-Fi和BLE两种协议的设备，实现了稳定的数据采集、协议转换和消息路由，跨协议联动功能运行正常，Web管理界面提供了便捷的设备管理和规则配置能力。研究揭示出并发接入时协议转换引擎的性能瓶颈，10设备并发时的时延较2设备场景增加约63%。本研究验证了采用统一数据模型和本地规则引擎实现边缘侧设备高效协同的有效路径。')
    
    # 6.2
    add_heading_custom(doc, '6.2 研究不足', level=2)
    add_heading_custom(doc, '(1) 研究反思', level=3)
    add_paragraph_custom(doc, '受限于实验设备和测试条件，本研究未能在大规模（如超过100台设备）并发场景下进行极限压力测试。在协议支持的广度方面，本研究仅对Wi-Fi和BLE两种协议进行了深度适配，Zigbee、Thread、Matter等协议的适配工作尚未开展。在安全性设计方面，设备身份认证采用简单的白名单机制，数据传输采用明文MQTT协议，未启用TLS加密。')
    
    # 6.3
    add_heading_custom(doc, '6.3 后续展望', level=2)
    add_heading_custom(doc, '(1) 展望建议', level=3)
    add_paragraph_custom(doc, '在智能化能力升级方面，可在网关中引入轻量级TinyML模型，通过对设备行为数据的边缘侧学习，实现从被动联动到主动智能预测的升级。在协议生态扩展方面，可在软件层面抽象出通用的协议驱动插件接口，实现更广范围协议的热插拔式兼容。在安全性强化方面，可引入基于X.509证书的双向TLS认证机制，实现基于角色的访问控制。在工程化部署方面，可探索容器化技术在网关软件交付中的应用，开发移动端APP作为Web后台的补充。')
    
    # ========== 参考文献 ==========
    doc.add_page_break()
    add_heading_custom(doc, '参考文献', level=1)
    
    references = [
        '[1] 杨飞宇. 面向智能家居的边缘智能网关的研究与设计[D]. 重庆理工大学, 2023.',
        '[2] 吴磊, 朱杰. 基于家庭多协议网关通信协议的设计与实现[J]. 计算机技术与发展, 2017, 27(9): 150-154.',
        '[3] 刘廷. 基于有线、无线通信技术的物联网智能家居系统[J]. 长江信息通信, 2022, 35(05): 186-188.',
        '[4] 葛悦涛, 尹晓桐. 边缘计算的发展趋势综述[J]. 无人系统技术, 2019, 2(02): 60-64.',
        '[5] 王哲. 边缘计算发展现状与趋势展望[J]. 自动化博览, 2021, 38(02): 22-29.',
        '[6] 杨帆. 浅谈智慧家居的移动物联网系统设计与应用[J]. 日用电器, 2021(11): 141-144.',
        '[7] 刘亮亮, 王兴, 王国庆, 等. 基于树莓派与MQTT的智能网关设计[J]. 机电工程技术, 2024, 53(08): 89-91+149.',
        '[8] 王猛. 多协议网关与智能家居的通信与控制[D]. 北方工业大学, 2017.',
        '[9] 方纪磊. 支持边缘计算任务的物联网网关系统的设计与实现[D]. 西安电子科技大学, 2023.',
        '[10] 王琦锋. 基于MQTT和ESP-NOW的智能家居监测与联动系统[D]. 宁夏大学, 2023.',
        '[11] Desai, P., Sheth, A., & Anantharam, P. Semantic gateway as a service architecture for IoT interoperability[C]. 2015 IEEE International Conference on Edge Computing (EDGE). IEEE, 2015: 265-272.',
        '[12] Perera, C., Zaslavsky, A., Christen, P., & Georgakopoulos, D. Sensing as a service model for smart cities supported by the internet of things[J]. Transactions on Emerging Telecommunications Technologies, 2014, 25(1): 81-93.',
        '[13] Li, Z., Zou, D., Xu, S., et al. VulPecker: an automated vulnerability detection system based on code similarity analysis[C]. Proceedings of the 32nd Annual Conference on Computer Security Applications. ACM, 2016: 201-213.',
        '[14] Ray, P. P. A survey on Internet of Things architectures[J]. Journal of King Saud University - Computer and Information Sciences, 2018, 30(3): 291-319.',
        '[15] Stojkovic, R., Gligorijevic, J., & Okanovic, D. Design and implementation of a smart home gateway[J]. Proceedings of the 2018 IEEE 26th International Scientific Conference Electronics (ET). IEEE, 2018: 1-4.',
    ]
    
    for ref in references:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(10.5)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = Cm(-0.74)
    
    print("第4-6章及参考文献添加完成")
    
    # 保存文档
    output_path = 'e:\\100_study\\120_Project\\CapstoneProject\\SmartHome_MultiProtocol_EdgeIntelligent_Gateway\\SH-MP-EG\\docs\\output\\面向智能家居的多协议边缘智能网关设计_论文_v2.docx'
    doc.save(output_path)
    print(f'完整论文文档已保存至: {output_path}')

if __name__ == '__main__':
    main()
