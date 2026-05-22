#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成完整论文Word文档（修复版）
包含正确的代码块、表格和图表占位符
"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_chinese_font(run, font_name='宋体', size=12, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold

def add_heading_custom(doc, text, level=1):
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    
    if level == 1:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(16)
        run.font.bold = True
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        heading.paragraph_format.space_before = Pt(24)
        heading.paragraph_format.space_after = Pt(18)
    elif level == 2:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(14)
        run.font.bold = True
        heading.paragraph_format.space_before = Pt(18)
        heading.paragraph_format.space_after = Pt(12)
    else:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(12)
        run.font.bold = True
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)
    
    heading.paragraph_format.line_spacing = 1.5
    return heading

def add_paragraph_custom(doc, text, first_line_indent=True, fontsize=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(fontsize)
    
    p.paragraph_format.line_spacing = 1.5
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_figure_placeholder(doc, fig_num, caption, note=None):
    """添加图片占位符 - 清晰格式"""
    # 占位符提示框
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(f'【在此处插入 图{fig_num}】')
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = None
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    
    # 图标题
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(f'图{fig_num} {caption}')
    set_chinese_font(run, '宋体', 10.5)
    
    if note:
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(f'({note})')
        set_chinese_font(run, '宋体', 9)
        p.paragraph_format.space_after = Pt(12)
    else:
        p.paragraph_format.space_after = Pt(12)

def add_table_caption(doc, caption):
    """添加表标题"""
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(caption)
    set_chinese_font(run, '宋体', 10.5, True)
    p.paragraph_format.space_after = Pt(6)

def add_table_from_data(doc, headers, rows):
    """添加表格"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # 表头
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(10)
                run.font.bold = True
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 数据行
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
            for p in row_cells[i].paragraphs:
                for run in p.runs:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(10)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    return table

def add_code_block(doc, code_lines, caption):
    """添加代码块 - 修复版，确保代码正确显示"""
    # 代码标题
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(caption)
    set_chinese_font(run, '宋体', 10.5, True)
    p.paragraph_format.space_after = Pt(6)
    
    # 创建代码框表格
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    
    # 设置表格宽度
    table.autofit = False
    table.allow_autofit = False
    
    cell = table.rows[0].cells[0]
    
    # 设置单元格宽度
    cell.width = Inches(6)
    
    # 设置灰色背景
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F0F0F0')
    shading_elm.set(qn('w:val'), 'clear')
    cell._element.get_or_add_tcPr().append(shading_elm)
    
    # 清空默认段落
    cell.text = ''
    
    # 添加代码行
    for line in code_lines:
        p = cell.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.left_indent = Cm(0.3)
    
    # 设置单元格边距
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for edge in ['top', 'left', 'bottom', 'right']:
        edge_elm = OxmlElement(f'w:{edge}')
        edge_elm.set(qn('w:w'), '100')
        edge_elm.set(qn('w:type'), 'dxa')
        tcMar.append(edge_elm)
    tcPr.append(tcMar)
    
    doc.add_paragraph()

# ========== 开始创建文档 ==========
doc = Document()

# 设置页面
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# 封面
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run('长沙学院')
set_chinese_font(run, '黑体', 26, True)
p.paragraph_format.space_after = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run('本科生毕业论文')
set_chinese_font(run, '黑体', 22, True)
p.paragraph_format.space_after = Pt(60)

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run('面向智能家居的多协议边缘智能网关设计')
set_chinese_font(run, '黑体', 18, True)
p.paragraph_format.space_after = Pt(80)

info_items = [
    ('学院', '计算机科学与工程学院'),
    ('专业', '物联网工程'),
    ('学生姓名', '黄志宇'),
    ('学号', 'B20220305131'),
    ('班级', '22物联01'),
    ('校内指导教师', '杨军'),
    ('职称', '高级工程师'),
]

for label, value in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(f'{label}：{value}')
    set_chinese_font(run, '宋体', 14)
    p.paragraph_format.line_spacing = 2.0

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run('2026年5月')
set_chinese_font(run, '宋体', 14)

doc.add_page_break()

# 摘要
add_heading_custom(doc, '摘  要', level=1)

abstract_paras = [
    '随着智能家居与物联网技术的迅猛发展，多协议设备的互联互通已成为制约用户体验提升的瓶颈[1]，传统的单一协议或云端集中处理方案难以兼顾低时延、高兼容性与本地自治需求。本文旨在设计并实现一款面向智能家居场景的多协议边缘智能网关，以解决异构设备接入困难与协同效率低下的问题。',
    
    '本研究采用系统设计与实验验证相结合的方法，基于树莓派和ESP32等硬件平台，设计了一套支持Wi-Fi、BLE与MQTT协议的软件系统[2]，并实现了协议适配、消息路由及本地联动等功能。系统采用四层分层架构设计，参照ETSI MEC参考架构[3]，实现了边缘侧数据处理和本地自治决策。',
    
    '通过搭建真实测试环境，对网关的功能完整性、传输时延及断网恢复能力进行了全面测试，结果表明：Wi-Fi传输时延平均值45.2ms，BLE时延平均值78.6ms，跨节点联动响应时延平均136.2ms[4]，均满足设计指标要求。系统在7×24小时长时运行测试中保持稳定运行，断网恢复能力测试补传成功率达到100%。',
    
    '本研究形成的软硬件参考架构能为嵌入式工程师、智能家居系统集成商提供具体、可复现的设计方案[5]。该方案在网关侧实现协议转换与本地联动规则引擎，本地控制响应时间控制在500毫秒以内，MQTT端到端传输延迟不超过200毫秒，在网络中断场景下依靠本地规则引擎仍能维持智能联动功能的正常运行[6]。',
    
    '本研究期望能为边缘计算在消费级物联网中的落地提供低成本、高效率的参考方案，推动智能家居系统向更开放、更自主的方向发展。'
]

for para in abstract_paras:
    add_paragraph_custom(doc, para)

p = doc.add_paragraph()
run = p.add_run('关键词：')
set_chinese_font(run, '黑体', 12, True)
run = p.add_run('边缘智能网关；多协议互联；智能家居；MQTT；协议转换')
set_chinese_font(run, '宋体', 12, True)
p.paragraph_format.line_spacing = 1.5

doc.add_page_break()

# ABSTRACT
add_heading_custom(doc, 'ABSTRACT', level=1)

abstract_en_paras = [
    'With the rapid development of smart home and Internet of Things technology, the interconnection of multi-protocol devices has become a bottleneck restricting the improvement of user experience. Traditional single-protocol or cloud-centric processing solutions are difficult to balance low latency, high compatibility, and local autonomy requirements. This paper aims to design and implement a multi-protocol edge intelligent gateway for smart home scenarios to solve the problems of heterogeneous device access difficulties and low collaboration efficiency.',
    
    'This research adopts a method combining system design and experimental verification, and designs a software system supporting Wi-Fi, BLE and MQTT protocols based on hardware platforms such as Raspberry Pi and ESP32, and implements protocol adaptation, message routing and local linkage functions. The system adopts a four-layer hierarchical architecture design, referring to the ETSI MEC reference architecture, and realizes edge-side data processing and local autonomous decision-making.',
    
    'By building a real test environment, the functional integrity, transmission delay and network disconnection recovery capability of the gateway were comprehensively tested. The results show that the average Wi-Fi transmission delay is 45.2ms, the BLE delay is 78.6ms, and the cross-node linkage response delay is 136.2ms on average, all meeting the design index requirements. The system maintained stable operation during the 7×24-hour long-term running test, and the relay success rate reached 100% in the network disconnection recovery capability test.',
    
    'This study expects to provide a low-cost and high-efficiency reference scheme for the implementation of edge computing in consumer-grade IoT, and promote the development of smart home systems towards a more open and autonomous direction.'
]

for para in abstract_en_paras:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(para)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

p = doc.add_paragraph()
run = p.add_run('Keywords: ')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run.font.bold = True
run = p.add_run('Edge Intelligent Gateway; Multi-protocol Interconnection; Smart Home; MQTT; Protocol Conversion')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run.font.bold = True
p.paragraph_format.line_spacing = 1.5

doc.add_page_break()

# 目录
add_heading_custom(doc, '目  录', level=1)

toc_items = [
    '摘  要................................................................................ I',
    'ABSTRACT............................................................................. II',
    '第1章 绪论.............................................................................. 1',
    '  1.1 研究背景与意义.................................................................. 1',
    '  1.2 国内外研究概况.................................................................. 3',
    '  1.3 研究内容........................................................................ 6',
    '  1.4 研究方法........................................................................ 6',
    '第2章 智能家居无线通信协议与关键技术分析............................................... 8',
    '  2.1 智能家居主流无线通信协议比较分析.................................................. 8',
    '  2.2 本系统支持协议与技术选型依据.................................................... 11',
    '  2.3 多协议共存与边缘侧协同可行性分析.................................................. 14',
    '第3章 多协议边缘网关总体方案与开发平台................................................ 16',
    '  3.1 网关总体方案设计............................................................... 16',
    '  3.2 硬件平台设计与实现............................................................. 18',
    '  3.3 软件平台搭建与运行环境配置..................................................... 21',
    '第4章 多协议边缘网关软件系统设计与实现................................................ 24',
    '  4.1 软件系统总体设计............................................................... 24',
    '  4.2 统一数据传输模型与主题规范设计................................................. 26',
    '  4.3 协议适配与数据收发模块设计实现................................................. 28',
    '  4.4 协议转换与消息路由模块设计实现................................................. 31',
    '  4.5 本地自治联动与离线补传机制实现................................................. 33',
    '  4.6 设备管理与状态管理功能实现..................................................... 36',
    '第5章 多协议边缘网关系统测试与验证................................................... 39',
    '  5.1 测试环境与测试方案............................................................. 39',
    '  5.2 系统功能测试与结果分析......................................................... 41',
    '  5.3 系统性能测试与结果分析......................................................... 43',
    '  5.4 系统稳定性测试与结果分析....................................................... 46',
    '  5.5 测试结果对比与达标分析......................................................... 48',
    '第6章 总结与展望...................................................................... 50',
    '  6.1 研究总结....................................................................... 50',
    '  6.2 研究不足....................................................................... 51',
    '  6.3 后续展望....................................................................... 52',
    '参考文献............................................................................. 54',
    '致  谢................................................................................. 55',
]

for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    set_chinese_font(run, '宋体', 12)
    p.paragraph_format.line_spacing = 1.5

print("封面、摘要、目录部分完成")

# 保存第一部分
doc.save('e:\\100_study\\120_Project\\CapstoneProject\\SmartHome_MultiProtocol_EdgeIntelligent_Gateway\\SH-MP-EG\\docs\\output\\论文_修复版_部分1.docx')
print("部分1保存完成")
